import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
from tkcalendar import Calendar
import os
import webbrowser
from PIL import Image
from PIL import ImageTk
import re
from datetime import datetime
from datetime import date
import calendar
import math
import openpyxl
from openpyxl.styles import Font
import string
from docx import Document
from docx.shared import Pt
import collections
import pymysql
import dbx_aws
import formatting
from decimal import Decimal, InvalidOperation
from win32api import GetSystemMetrics
import importlib
import threading
import time

screen_width = GetSystemMetrics(0)
screen_height = GetSystemMetrics(1)
width = int(screen_width/1.2)
height = int(screen_height/1.2)
command_width = int(.33*width)
table_width = int(.66*width)
x = (screen_width/2) - (width/2)
y = ((screen_height/2) - (height/2)) - 50

# Calendar button image
cal_image = Image.open("calendar.png")
cal_image = cal_image.resize((22,22))

def db_update():
    conn = dbx_aws.db_conn()
    cur = conn.cursor()
    # Updates inventory database values for specific columns.
    cur.execute("""UPDATE bottles
                      SET amount=0
                    WHERE amount<0
                """)
    cur.execute("""UPDATE samples
                      SET amount=0
                    WHERE amount<0
                """)
    cur.execute("""SELECT COUNT(*)
                     FROM barrels
                """)
    barrel_count = str(cur.fetchone()[0])
    cur.execute("UPDATE barrel_count " +
                   "SET full_amount=" + barrel_count)
    cur.execute("""UPDATE barrel_count
                      SET total=(full_amount + empty_amount) * price
                """)
    cur.execute("""UPDATE barrels
                   SET age=(DATEDIFF(NOW(), date_filled))
                """)
    cur.execute("""UPDATE estimated_cogs
                      SET total_per_bottle=format(
                           (raw_mat + energy + labor + error), 3)
                """)
    cur.execute("""UPDATE estimated_cogs
                    SET total_per_pg=format(
                         ((total_per_bottle*mult_fact) + bond_ins + storage), 3)
                """)
    cur.execute("""UPDATE raw_materials
                      SET total=(amount * price)
                """)
    cur.execute("""UPDATE bottles
                      SET total=(amount * price)
                """)
    cur.execute("""UPDATE samples
                      SET total=(amount * price)
                """)
    cur.execute("""UPDATE grain
                      SET total=(amount * price)
                """)
    conn.commit()
    conn.close()

def edit_db(sql_edit, sql_table, gui_table, view_fr, delete=False):
    # Updates the sql_table with the changes provided by sql_edit.
    # sql_edit is a tuple of length 2*(num of cols)
    conn = dbx_aws.db_conn()
    cur = conn.cursor()
    cur.execute("SELECT * FROM {}".format(sql_table))
    columns = [x[0] for x in cur.description]
    str_and = "=%s AND ".join(columns) + "=%s" # col1=? AND col2=? AND ..

    if delete == False:
        # Create 2 equal length lists to be formatted
        half = int(len(sql_edit) / 2) # List length will always be even
        l1,l2 = sql_edit[:half], sql_edit[half:]
        l1_formatted = formatting.reverse_format_data(l1, columns)
        l2_formatted = formatting.reverse_format_data(l2, columns)
        sql_edit = tuple(l1_formatted + l2_formatted)
        str_comma = "=%s, ".join(columns) + "=%s" # col1=?, col2=?, ...
        query = "UPDATE {} SET {} WHERE {}".format(sql_table, str_comma,
                                                   str_and)
        cur.execute(query, sql_edit)
    else:
        sql_edit = formatting.reverse_format_data(sql_edit, columns)
        query = "DELETE FROM {} WHERE {}".format(sql_table, str_and)
        cur.execute(query, sql_edit)

    conn.commit()
    cur.close()
    conn.close()
    db_update()

    try:
        view_fr.columns.set("All")
        view_fr.columns.event_generate("<<ComboboxSelected>>")
    except:
        pass
    try:
        view_products(sql_table, 'All', 'All', gui_table)
    except:
        pass


def monthly_reports_update():
    # Update 'monthly_reports' table with current month and inventory
    # values. Only select purchase orders from current month. Selects
    # pending purchase orders to be fulfilled next month.
    inv_tables = ['raw_materials', 'bottles', 'samples', 'pending_po', 'grain',
                  'purchase_orders',
    ]
    monthly_totals = collections.OrderedDict()
    cur_month = datetime.today()
    try:
        next_month = cur_month.replace(month=cur_month.month+1)
    except ValueError:
        if x.month == 12:
            next_month = cur_month.replace(year=cur_month.year+1, month=1)
    cur_month = cur_month.strftime("%Y-%m-01")
    next_month = next_month.strftime("%Y-%m-01")
    conn = dbx_aws.db_conn()
    cur = conn.cursor()
    cur.execute("SELECT product, price FROM bottles")
    prod_prices = {key:val for (key, val) in cur.fetchall()}
    for table in inv_tables:
        if table == 'pending_po':
            query = ("SELECT product, amount, total FROM {} WHERE po_date " +
                     "BETWEEN %s and %s").format(table)
            cur.execute(query, (cur_month, next_month))
            pending_info = cur.fetchall()
            pend_sale_amts = {} #Total sale amounts by product
            for (prod, amt) in [x[:2] for x in pending_info]:
                if prod in pend_sale_amts:
                    pend_sale_amts[prod] += int(amt)
                else:
                    pend_sale_amts[prod] = int(amt)
            pend_cogs_total = 0
            for prod in pend_sale_amts.keys():
                pend_cogs_total += float(pend_sale_amts[prod]
                                         * prod_prices[prod])
            pend_sale_total = sum([float(x[2]) for x in pending_info])
            monthly_totals['pending_sales'] = pend_sale_total
            monthly_totals['pending_cogs'] = -1*pend_cogs_total

        elif table == 'purchase_orders':
            query = ("SELECT product, amount, total FROM {} WHERE pu_date " +
                     "BETWEEN %s and %s").format(table)
            cur.execute(query, (cur_month, next_month))
            po_info = cur.fetchall()
            po_sale_amts = {}   #Total sales amount by product
            for (prod, amt) in [x[:2] for x in po_info]:
                if prod in po_sale_amts.keys():
                    po_sale_amts[prod] += int(amt)
                else:
                    po_sale_amts[prod] = int(amt)
            try:
                po_cogs_total = 0
                for prod in po_sale_amts.keys():
                    po_cogs_total += float(po_sale_amts[prod]
                                           * prod_prices[prod])
            except:
                po_cogs_total = 0
            po_sale_total = sum([float(x[2]) for x in po_info])
            monthly_totals['purchase_order_sales'] = po_sale_total
            monthly_totals['purchase_order_cogs'] = -1*po_cogs_total

        else:
            cur.execute("SELECT total FROM {}".format(table))
            total_vals = [x[0] for x in cur.fetchall()]
            total = sum([float(x) for x in total_vals])
            monthly_totals[table] = total

    cur.execute("SELECT * FROM barrels")
    barrel_vals = cur.fetchall()
    cur.execute("SELECT total_per_pg FROM estimated_cogs")
    est_cogs = [x[0] for x in cur.fetchall()]
    whisk_cogs = float(est_cogs[0])
    rum_cogs = float(est_cogs[1])
    whisk_total = 0
    rum_total = 0
    for barrel in barrel_vals:
        if barrel[1] == 'Rum':
            rum_total += float(barrel[3]) * rum_cogs
        else:
            whisk_total += float(barrel[3]) * whisk_cogs
    monthly_totals['barreled_rum'] = ("%.2f" % rum_total)
    monthly_totals['barreled_whiskey'] = ("%.2f" % whisk_total)
    cur.execute("SELECT total FROM barrel_count")
    barr_total = cur.fetchone()[0]
    monthly_totals['barrels'] = barr_total

    for key, value in monthly_totals.items():
        value = ("%.2f" % float(value))
        query = ("INSERT INTO monthly_reports VALUES (%s, %s, %s) " +
                 "ON DUPLICATE KEY UPDATE total=%s")
        cur.execute(query, (cur_month, key, value, value))
    conn.commit()
    cur.close()
    conn.close()

def create_excel_inv():
    # Populate inventory_template.xlsx with inventory values and save as
    # new workbook.
    # ('inv table', 'total column index')
    total_cols = (
        ('raw_materials', 5),
        ('production', None),
        ('in_progress', None),
        ('bottles', 6),
        ('samples', 6),
        ('grain', 6),
        ('mashes', None),
        ('grain_log', None),
        ('barrels', None),
        ('empty_barrels', None),
        ('purchase_orders', 7),
        ('pending_po', 7),
        ('employee_transactions', None),
    )
    total_cols = collections.OrderedDict(total_cols) # Maintain order
    excel_file = (os.getcwd() + "/temp_files/inventory_template.xlsx")
    wb = openpyxl.load_workbook(excel_file)
    sheets = wb.sheetnames
    for sheet, (inv, tot_col) in zip(sheets, total_cols.items()):
        active_sheet = wb[sheet]
        conn = dbx_aws.db_conn()
        query = "SELECT * FROM {}".format(inv)
        inv_values = dbx_aws.db_exec(conn, query)
        inv_values = formatting.format_data(inv_values.data, inv_values.columns)
        if len(inv_values) > 0:
            for (indx, row) in enumerate(inv_values, 2):
                try:
                    row = list(row)
                    active_sheet.append(row)
                    active_sheet.cell(indx, tot_col).number_format = '$#,##0.00'
                except TypeError:
                    pass
            last_row = len(inv_values) + 1
            last_col = len(inv_values[0]) - 1
            rows = range(1, last_row + 1)
            columns = string.ascii_uppercase[:last_col + 1]
            last_col = string.ascii_uppercase[last_col]
            # Format text to justify centrally
            for row in rows:
                for col in columns:
                    cell = col + str(row)
                    active_sheet[cell].alignment = (
                        openpyxl.styles.Alignment(horizontal='center'))
            # Format excel table
            tbl_ref = "A1:{}{}".format(str(last_col), str(last_row))
            tbl = openpyxl.worksheet.table.Table(displayName=inv, ref=tbl_ref)
            style = openpyxl.worksheet.table.TableStyleInfo(
                name="TableStyleMedium9", showFirstColumn=False,
                showLastColumn=False, showRowStripes=True)
            tbl.tableStyleInfo = style
            active_sheet.add_table(tbl)
    # Save and open excel file
    cur_month = datetime.now().strftime("%Y-%m")
    temp_excel_file = os.path.abspath(os.getcwd() + r"\temp_files\temp_monthly.xlsx")
    db_file_name = "{}.xlsx".format(cur_month)
    wb.save(temp_excel_file)
    dbx_aws.Dropbox_Excel_Upload('/ADCo Team Folder/Operations - ALL/Montgomery/Luc File/Inventory/MONTHLY INVENTORIES',
                                 temp_excel_file, db_file_name)
    open_ques = messagebox.askquestion(
        "Open the Excel File?",
        "Would you like to open a temporary version of the Monthly Inventory " +
        "file in Excel? This will allow you to print it now. \n" +
        "***THIS IS NOT THE ACTUAL FILE. THAT FILE IS IN THE DROPBOX FOLDER.")
    if open_ques == "yes":
        try:
            os.system('start EXCEL.EXE ' + temp_excel_file)
        except:
            messagebox.showerror(
                'Program Error',
                'There was an error opening Excel.'
            )
    else:
        pass


def view_widget(window, widget, location, sql_table, column, item,
                gui_table):
    # Removes current packed widgets from window frame and replaces with
    # new widget chosen.
    for widg in window.pack_slaves():
        widg.pack_forget()
    widget.pack(side=location, fill='both', expand=1)
    if gui_table:
        view_products(sql_table, column, item, gui_table)


def view_products(sql_table, column, item, gui_table):
    # Fetches info from sql_table based on an item filter.  Returns
    # information into the current gui_table.  Configures even-numbered
    # rows to have a grey background.
    conn = dbx_aws.db_conn()
    if column == "All":
        query = "SELECT * FROM {}".format(sql_table)
        data = dbx_aws.db_exec(conn, query)

    elif column == "barrel_no":
        query = "SELECT * FROM {} WHERE {} LIKE \'{}%\'".format(sql_table,
                                                                column,
                                                                item[:2])
        data = dbx_aws.db_exec(conn, query)
    elif 'date' in column:
        cur_year = datetime.strptime(item, "%Y")
        next_year = cur_year.replace(year=cur_year.year+1)
        query = "SELECT * FROM {} WHERE {} BETWEEN %s and %s".format(sql_table,
                                                                     column)
        data = dbx_aws.db_exec(conn, query, True, cur_year.strftime("%Y-01-01"),
                               next_year.strftime("%Y-01-01"))
    elif column == "age":
        cur_age = int(item[0]) * 365  # Number of years in days
        next_age = (int(item[0]) + 1) * 365 # Number of years+1 in days
        query = "SELECT * FROM {} WHERE {} BETWEEN %s and %s".format(sql_table,
                                                                     column)
        data = dbx_aws.db_exec(conn, query, True, cur_age, next_age)
    elif column == "product":
        query = "SELECT * FROM {} WHERE {} LIKE '{}%'".format(sql_table,
                                                                column,
                                                                item)
        data = dbx_aws.db_exec(conn, query)
    else:
        query = "SELECT * FROM {} WHERE {}='{}'".format(sql_table,
                                                           column,
                                                           item)
        data = dbx_aws.db_exec(conn, query)
    data.data = formatting.format_data(data.data, data.columns)

    for item in gui_table.get_children(): # Remove old table values
        gui_table.delete(item)

    for (index, row) in enumerate(data.data, 1): # Created striped rows
        if (index % 2 == 0):
            tag = 'even'
        else:
            tag = 'odd'
        gui_table.insert("", 'end', values=row, tags=(tag,))
    gui_table.tag_configure('even', background='#E8E8E8')

    if sql_table == 'barrels':
        from __main__ import gui
        gui.barr_count_fr.barr_update(first=1)

def file_view(folder, master):
    # Displays a toplevel window populated by clickable links to files
    # within the given folder.
    labels_window = tk.Toplevel(master)
    files = os.listdir(os.getcwd() + "\\" + folder)
    window_height = 0
    for file in files:
        mo = fileRegex.search(file)
        file_name = mo.group(1).replace("_", " ")
        file_label = SheetLabel(
            master=labels_window, text=file_name,
            file_location="{}\\{}\\{}".format(os.getcwd(), folder, file)
        )
        file_label.pack(padx=10, pady=5, anchor='w')
        window_height += 38
    labels_window.title(folder.replace("_", " ").title())
    labels_window.focus()
    x = (screen_width/2) - (250)
    y = (screen_height/2) - (250)
    labels_window.geometry("%dx%d+%d+%d" % (300, window_height, x, y))
    labels_window.resizable(0,0)


def selection_check(sql_table, gui_table, view_fr, window, barr_count_fr,
                    edit=True, delete=False, empty=False):
    # Checks to see if a gui_table selection has been made and returns
    # the respective action based on the gui_table.
    item_values = gui_table.item(gui_table.selection())['values']
    if item_values:
        if delete == True:
            del_ques = messagebox.askquestion(
            "Delete Current Selection?",
            "Are you sure you want to continue? Confirming will delete the " +
            "current selection from the inventory. This information will not " +
            "be able to be recovered.")
            if del_ques == 'yes':
                edit_db(list(item_values), sql_table, gui_table, view_fr,
                        delete=True)
            else:
                return
        elif empty == True:
            EmptyBarrelView(window, item_values)
        # Open po excel file
        elif (sql_table == 'purchase_orders' and edit==False):
            po_num = item_values[8]
            excel_file = "/{}.xlsx".format(po_num)
            dbx_aws.Dropbox_Excel_Download(dbx_aws.po_folder, excel_file)
        elif (sql_table == 'in_progress'): # Finish in progress production
            FinishView(window, item_values)
        elif (sql_table == 'pending_po' and edit==False): # Fulfill po
            fulfill_pending(gui_table, view_fr)
        else: # Edit selection
            EditView(window, sql_table, gui_table, 2, view_fr, barr_count_fr)
    else:
        messagebox.showerror(
            "Selection Error",
            "Please select an inventory item.", parent=window)


def gui_table_sort(gui_table, column, reverse):
    # Sorts gui tables in ascending order based on the column header
    # clicked.  The next click upon the header will be in reverse order.
    l = [(gui_table.set(k, column), k) for k
         in gui_table.get_children()]
    if '$' in l[0][0]: # Check if column is 'total'
        l.sort(key=lambda tup: float(tup[0][1:].replace(",", "")),
               reverse=reverse)
    else:
        try:
            l.sort(key=lambda tup: float(tup[0].replace(",","")),
                   reverse=reverse)
        except ValueError:
            l.sort(key=lambda tup: tup[0], reverse=reverse)

    # Rearrange items in sorted positions.
    for index, (val, k) in enumerate(l):
        gui_table.move(k, '', index)
        gui_table.item(k, tags=())
        if index % 2 == 0:
            gui_table.item(k, tags=('even',))
    gui_table.tag_configure('even', background="#E8E8E8")

    # Reverse sort next time.
    gui_table.heading(
        column, text=column,
        command=lambda c=column: gui_table_sort(gui_table, c, not reverse)
    )


def cal_button(tplvl, date_entry, window):
    # Creates a toplevel window to provide a calendar date selection
    # tool.
    tplvl.top = tk.Toplevel(window)
    tplvl.cal = Calendar(tplvl.top, font="Arial 14", selectmode='day',
                         locale='en_US', cursor="hand2")
    tplvl.cal.pack(fill="both", expand=True)
    (HoverButton(tplvl.top, text="ok",
                 command=lambda: retrieve_date(tplvl, date_entry))
                 .pack())
    tplvl.top.focus()


def retrieve_date(tplvl, date_entry):
    # Updates the date-entry widget within the toplevel widget with
    # formatted date value.
    date_entry.config(state='normal')
    date_entry.delete(0 ,'end')
    date_entry.insert('end', tplvl.cal.selection_get().strftime("%Y-%m-%d"))
    date_entry.config(state="readonly")
    tplvl.top.destroy()


def confirm_po(info, purchase_orders, po_num):
    # Insert purchase order info into 'purchase_orders' table and create
    # excel file with purchase order info.
    year = datetime.now().year
    wb = openpyxl.load_workbook('temp_files/blank_po.xlsx')
    sheet = wb['Purchase Order']
    font = Font(name='Times New Roman', size=12)

    # Get shipment information entry-values into list and input them
    # into corresponding cells within the 'po' excel sheet.
    info_cells = ['A9', 'K9', 'A12', 'A15', 'I15']
    for entry,cell in zip(info, info_cells):
        sheet[cell] = entry
        sheet[cell].font = font
    # Get order values and input into table within purchase order excel
    # file.
    excel_columns = ["A","B","D","J","M"]
    excel_rows = range(18, 36)
    index = 0
    total_po = 0
    for i in excel_rows:
        for j,k in zip(excel_columns, range(0,5)):
            cell = j + str(i)
            try:
                sheet[cell] = (purchase_orders[index][k])
                sheet[cell].font = font
                if k == 4:
                    try:
                        total_po += float(
                            purchase_orders[index][k].replace(',',''))
                    except AttributeError: # If num is in decimal format
                        total_po += purchase_orders[index][k]
            except IndexError:
                sheet[cell] = ""
        index += 1
    total_po = Decimal(total_po).quantize(Decimal('.01'),
                                          rounding=ROUND_HALF_UP)
    sheet['M36'] = total_po
    sheet['M36'].number_format = '$#,##0.00'

    # Add purchase orders to 'purchase_orders' table and update
    # inventory.
    conn = dbx_aws.db_conn()
    cur = conn.cursor()
    for po_list in (x for x in purchase_orders if all(x)):
        query = "INSERT INTO purchase_orders VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)"
        query_vals = (
            info[3], info[4], po_list[2], po_list[0], po_list[1], po_list[3],
            po_list[4], info[2], po_num,
        )
        cur.execute(query, query_vals)
        if po_list[1] == "Cases":
            query = "UPDATE bottles SET amount=(amount - %s) WHERE product=%s"
            cur.execute(query, (po_list[0], po_list[2]))
        else:
            query = "UPDATE samples SET amount=(amount - %s) WHERE product=%s"
            cur.execute(query, (po_list[0], po_list[2]))
    conn.commit()
    cur.close()
    conn.close()

    temp_file = os.path.abspath(os.getcwd() + r'\temp_files\temp_po.xlsx')
    file_name = "{}.xlsx".format(po_num)
    wb.save(temp_file)
    dbx_aws.Dropbox_Excel_Upload('/ADCo Team Folder/Operations - ALL/Montgomery/Luc File/Inventory/PURCHASE ORDERS',
                                 temp_file, file_name)

    open_ques = messagebox.askquestion(
        "Open the PO Excel File?",
        "Would you like to open a temporary version of the Purchase Order " +
        "file in Excel? This will allow you to print it now.")
    if open_ques == "yes":
        try:
            os.system('start EXCEL.EXE ' + temp_file)
        except:
            messagebox.showerror(
                "Program Error",
                "There was an error opening Excel.")
    else:
        pass

    db_update()
    from __main__ import gui
    gui.po_vfr.columns.set("All")
    gui.po_vfr.columns.event_generate("<<ComboboxSelected>>")

def fulfill_pending(gui_table, view_fr):
    # Remove pending purchase order from 'pending_po' table and input
    # into 'purchase_orders'. Create an excel file with info.
    po_num = gui_table.item(gui_table.selection())['values'][8]
    conn = dbx_aws.db_conn()
    query = "SELECT * FROM pending_po WHERE po_number=\'{}\'".format(po_num)
    po_vals = dbx_aws.db_exec(conn, query).data
    po_1 = po_vals[0] # First PO to retrieve info from
    po_info = [
        "Montgomery", # Departure location (always Montgomery)
        po_num, # PO Number
        po_1[7], # Destination
        po_1[0], # PO Date
        po_1[1] # Pick up Date
    ]
    comp_pos = [
        [
         x[3], # amount
         x[4], # unit
         x[2], # product
         x[5], # price
         x[6], # total
        ] for x in po_vals
    ]
    confirm_po(po_info, comp_pos, po_num)
    # Reformat data to avoid issues with edit_db
    po_cols = [
        'po_date', 'pu_date', 'product', 'amount', 'unit',
        'price', 'total', 'destination', 'po_number'
    ]
    po_vals = formatting.format_data(po_vals, po_cols)
    for po in po_vals: # Remove pending PO values from SQL
        edit_db(po, 'pending_po', gui_table, view_fr, delete=True)


class HoverButton(tk.Button):
    # Button widget with mouse-over color and cursor changes.
    def __init__(self, master, **kw):
        tk.Button.__init__(self, master=master, cursor="hand2", **kw)
        self.defaultBackground = self["background"]
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
        self.bind("<Button-1>", self.on_click)

    def on_enter(self, event):
        self['background'] = 'gray70'

    def on_leave(self, event):
        self['background'] = self.defaultBackground

    def on_click(self, event):
        self['cursor'] = 'wait'
        time.sleep(.1)
        self['cursor'] = 'hand2'


class AddView(tk.Toplevel):
    # A toplevel widget with labels corresponding to sql table
    # columns and entry widgets to insert data into the sql table.
    def __init__(self, master, sql_table, gui_table, entry_col, view_fr,
                 barr_count_fr):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.entry_col = entry_col
        self.view_fr = view_fr
        self.barr_count_fr = barr_count_fr
        self.tplvl_title = self.sql_table.replace("_"," ").title()
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        tk.Toplevel.__init__(self, master=master)

        self.title_frame = tk.Frame(self)
        (tk.Label(self.title_frame,
                  text="Add Product to {} Inventory".format(self.tplvl_title),
                  font="Arial 10 bold"
                 ).pack()
        )
        self.title_frame.grid(row=0, column=0, columnspan=2, pady=5)

        # Create labels and entries based on gui_table column headers.
        for (index, description) in enumerate(self.gui_table.columns, 1):
            if (description.lower() != 'type'):
                tk.Label(self, text="{}:".format(description)).grid(row=index,
                                                                    column=0)
                if description.lower() == 'total':
                    self.total_text = tk.StringVar()
                    self.total_entry = tk.Entry(
                        self, textvariable=self.total_text
                    )
                    self.total_entry.config(state="readonly")
                    self.total_entry.grid(row=index, column=self.entry_col)
                elif (description.lower().find('age') != -1):
                    self.age_text = tk.StringVar()
                    self.age_entry = tk.Entry(self, textvariable=self.age_text)
                    self.age_entry.config(state="readonly")
                    self.age_entry.grid(row=index, column=self.entry_col)
                elif (description.lower() == 'price'):
                    (tk.Entry(
                        self, validate='key',
                        validatecommand=(self.register(valid_dec),
                                         '%S', '%s', '%d')
                             ).grid(row=index, column=self.entry_col)
                    )
                elif (description.lower() == 'amount'):
                    (tk.Entry(
                        self, validate='key',
                        validatecommand=(self.register(valid_dig), '%S','%d')
                             ).grid(row=index, column=self.entry_col)
                    )
                else:
                    tk.Entry(self).grid(row=index, column=self.entry_col)
                if (description.lower().find('date') != -1):
                    self.date_index = index
                    self.date_entry = self.grid_slaves(row=self.date_index,
                                                       column=self.entry_col)[0]
                    self.date_entry.config(state="readonly")
                    self.cal_photo = ImageTk.PhotoImage(cal_image)
                    self.cal_link = HoverButton(
                        self, image=self.cal_photo,
                        command=lambda: cal_button(self, self.date_entry,
                                                   self.master)
                    )
                    self.cal_link.image = self.cal_photo
                    self.cal_link.grid(row=index, column=self.entry_col+1)
                elif ((description.lower().find('total') != -1) or
                      (description.lower().find('age') != -1)):
                    self.labels = [x for x
                                   in reversed(self.grid_slaves(column=0))
                                   if (x.winfo_class() == 'Label')]
                    for entry in self.labels:
                        if entry.cget("text").lower().find("amount") != -1:
                            self.amount_row = entry.grid_info()['row']
                            self.amount_entry = self.grid_slaves(
                                row=self.amount_row,
                                column=self.entry_col)[0]
                        if entry.cget("text").lower().find("price") != -1:
                            self.price_row = entry.grid_info()['row']
                            self.price_entry = self.grid_slaves(
                                row=self.price_row,
                                column=self.entry_col)[0]
                    self.total_after()
            else: # Type option entry.
                tk.Label(self, text="{}:".format(description)).grid(row=index,
                                                                    column=0)
                self.options = ttk.Combobox(self,
                                            values=type_options[sql_table])
                self.options.set(type_options[sql_table][0])
                self.options.config(width=16, background="white",
                                    justify='center', state='readonly')
                self.options.grid(row=index, column=self.entry_col)

        self.grid_size = self.grid_size()[1]
        self.button_frame = tk.Frame(self)
        (HoverButton(self.button_frame, text="Add Item", width=10,
                     command=lambda: busy_cursor(self, self.confirm)
                    ).pack(side='left', padx=5, pady=5)
        )
        (HoverButton(self.button_frame, text="Cancel", width=10,
                     command=lambda : self.destroy()
                    ).pack(side='left', padx=5, pady=5)
        )
        self.button_frame.grid(row=self.grid_size+1, column=0, columnspan=2)

        self.title("Add to {}".format(self.tplvl_title))
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        # Work through AddView toplevel to find entry widgets and
        # extract these values to be inserted into the sql table.
        # Uses db_update() to update certain column values afterwards
        # and view_products() to display the updated gui table.
        self.additions = [] # List to be populated by entries
        self.entries = [x for x
                        in reversed(self.grid_slaves())
                        if (x.winfo_class() == 'Entry'
                        or x.winfo_class() == 'TCombobox')]

        for entry in self.entries:
            if entry.get():
                self.additions.append(entry.get())
            else:
                messagebox.showerror("Input Error",
                                     "At least one input is blank, " +
                                     "please try again.", parent=self)
                return
        self.additions = tuple(self.additions)
        self.db_add(self.sql_table, self.additions)

        if self.sql_table == 'barrels':
            db_update()

        try:
            self.view_fr.columns.set("All")
            self.view_fr.columns.event_generate("<<ComboboxSelected>>")
        except:
            view_products(self.sql_table, 'All', 'All', self.gui_table)

        self.master.config(cursor='arrow')
        self.master.update()
        self.destroy()

    def db_add(self, table, data):
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        self.str1 = "({}%s)".format("%s,"*(len(data) - 1)) #(%s,%s,..,%s)
        self.query = "INSERT INTO {} VALUES {}".format(table, self.str1)
        self.cur.execute(self.query, data)
        self.conn.commit()
        self.cur.close()
        self.conn.close()

    def total_after(self):
        # Widget after-function to update total and age entry values.
        # Repeats every 150ms.
        def total_update():
            # Tries to update total and age entry values.
            # Raises:
            # AttributeError: if price_entry, amount_entry, date_entry
            # don't exist
            # ValueError: if price_entry, amount_entry, date_entry
            # values are currently empty
            try: # Update total entry
                self.price_num = self.price_entry.get()
                self.amount_num = self.amount_entry.get()
                self.total_value = round((Decimal(self.amount_num)
                                           * Decimal(self.price_num)),
                                           2)
                self.total_text.set(self.total_value)
                return
            except (AttributeError, ValueError, InvalidOperation):
                pass
            try:    # Update age entry.
                self.date_value = datetime.strptime(self.date_entry.get(),
                                                    '%Y-%m-%d')
                self.date_diff = datetime.now() - self.date_value
                self.age_text.set(self.date_diff.days)
            except (AttributeError, ValueError):
                pass
        total_update()
        self.after(150, self.total_after)


class EditView(AddView):
    # A toplevel widget with labels corresponding to sql table
    # columns and entry widgets to update data in sql table.
    def __init__(self, master, sql_table, gui_table, entry_col, view_fr,
                 barr_count_fr):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.entry_col = entry_col # Column location for entry widgets.
        self.view_fr = view_fr
        self.selection = self.gui_table.selection()
        self.item_values = self.gui_table.item(self.selection)['values']
        self.tplvl_title = self.sql_table.replace("_"," ").title()
        AddView.__init__(self, master, sql_table, gui_table, entry_col,
                          view_fr, barr_count_fr)
        self.title("Edit {}".format(self.tplvl_title))

        self.title_frame.destroy() # Remove AddView title frame.
        self.title_frame = tk.Frame(self)
        (tk.Label(
            self.title_frame,
            text="Edit Product in {} Inventory".format(self.tplvl_title),
            font="Arial 10 bold"
                 ).pack()
        )
        self.title_frame.grid(row=0, column=0, columnspan=3)

        # Create toplevel labels.
        for (index, description) in enumerate(self.gui_table.columns):
            (tk.Label(
                self, text=self.item_values[index], foreground='blue'
                     ).grid(row=index+1, column=1)
            )

        self.button_frame.destroy()
        self.button_frame = tk.Frame(self)
        (HoverButton(
            self.button_frame, text="Confirm", command=self.confirm
                    ).pack(side='left', padx=5, pady=5)
        )
        (HoverButton(self.button_frame, text="Cancel",
                     command=lambda: self.destroy()
                    ).pack(side='left', padx=5, pady=5)
        )
        self.button_frame.grid(row=self.grid_size+1, column=0, columnspan=3)

    def confirm(self):
        # Work through Edit_View toplevel to find entry widgets and
        # extract these values to be updated in the given sql table.
        # Uses db_update() to update certain column values afterwards
        # and view_products() to display the updated gui table.
        self.changes = [] # List where updated entries will exist
        self.edit_entries = [x for x
                             in reversed(self.grid_slaves())
                             if (x.winfo_class() == 'Entry'
                             or x.winfo_class() == 'TCombobox')]
        for entry in self.edit_entries:
            if entry.get():
                self.changes.append(entry.get())
            else:
                messagebox.showerror("Input Error",
                                     "At least one input is blank, please try" +
                                     " again.", parent=self)
                return
        self.current_values = [x.cget('text') for x
                               in reversed(self.grid_slaves(column=1))
                               if x.winfo_class() == 'Label'
                              ]
        self.sql_edit = tuple(self.changes + self.current_values)
        edit_db(self.sql_edit, self.sql_table, self.gui_table, self.view_fr)
        self.destroy()


class ProductionView(tk.Toplevel):
    # Toplevel used to register production.  Subtracts values from raw
    # materials when used.  Handles unfinished products by placing in
    # 'in_progress' table to be finished later.
    def __init__(self, master, sql_table, gui_table):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        tk.Toplevel.__init__(self, master=master)

        self.title_frame = tk.Frame(self)
        (tk.Label(
            self.title_frame, text="Production", font="Arial 10 bold"
                 ).pack()
        )
        self.title_frame.grid(row=0, column=0, columnspan=3, pady=5)

        self.product_frame = tk.Frame(self)
        tk.Label(self.product_frame, text="Total Bottles").grid(row=0, column=0)
        tk.Label(self.product_frame, text="Cases").grid(row=0, column=1)
        tk.Label(self.product_frame, text="Product").grid(row=0, column=2)
        (tk.Entry(
            self.product_frame, validate='key',
            validatecommand=(self.register(valid_dig),'%S','%d')
                 ).grid(row=1, column=0, padx=5)
        )
        (tk.Entry(
            self.product_frame, validate='key',
            validatecommand=(self.register(valid_dig),'%S','%d')
                 ).grid(row=1, column=1, padx=5)
        )
        self.conn = dbx_aws.db_conn()
        self.query = "SELECT product FROM bottles"
        self.product_rows = dbx_aws.db_exec(self.conn, self.query)
        # Extract product names from list of tuples
        self.product_rows = [x[0] for x in self.product_rows.data]
        self.products = ttk.Combobox(self.product_frame,
                                     values=self.product_rows)
        self.products.config(width=20, background="white", justify='center',
                             state='readonly')
        self.products.set(self.product_rows[0])
        self.products.grid(row=1, column=2, padx=5)
        self.product_frame.grid(row=1, column=0, columnspan=3)

        # Raw materials title frame.
        self.materials = tk.Frame(self)
        (tk.Label(
            self.materials, text="Materials Used", font="Arial 10 bold"
                 ).pack()
        )
        self.materials.grid(row=3, column=0, columnspan=3, pady=5)

        # Raw materials input frame.
        tk.Label(self, text="Type").grid(row=4, column=0, pady=2)
        tk.Label(self, text="Amount").grid(row=4, column=1, pady=2)
        tk.Label(self, text="Material").grid(row=4, column=2, pady=2)
        self.type_rows = type_options['raw_materials']

        # Create label, entry, option box for each type of raw material.
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        for (index, description) in enumerate(self.type_rows, 5):
            tk.Label(self, text="{}:".format(description)).grid(row=index,
                                                                column=0,
                                                                sticky='W')
            (tk.Entry(
                self, validate='key',
                validatecommand=(self.register(valid_dig),'%S','%d')
                     ).grid(row=index, column=1)
            )
            self.query = ("SELECT product FROM raw_materials " +
                          "WHERE type=\'{}\'").format(description)
            self.cur.execute(self.query)
            # Extract materials from list of tuples
            self.rows = [x[0] for x in self.cur.fetchall()]
            self.rows.append("None")
            self.opt_menu = ttk.Combobox(self, values=self.rows)
            self.opt_menu.config(width=20, background="white", justify='center',
                                 state='readonly')
            self.opt_menu.set(self.rows[0])
            self.opt_menu.grid(row=index, column=2, padx=5)
        self.cur.close()
        self.conn.close()

        # Finished product checkbox.
        self.grid_size = self.grid_size()[1]
        self.check_var = tk.IntVar()
        self.check_var.set(1)
        self.check_b = tk.Checkbutton(
            self, text="Are the products finished? (i.e. labeled)",
            variable=self.check_var, command=self.cbox_check)
        self.check_b.grid(row=self.grid_size+1, column=0, columnspan=3)

        # Samples input frame.
        self.samples_frame = tk.Frame(self)
        tk.Label(self.samples_frame, text="Samples").grid(row=0, column=0)
        self.samples_entry = tk.Entry(
            self.samples_frame, validate='key',
            validatecommand=(self.register(valid_dig),'%S','%d')
        )
        self.samples_entry.grid(row=0, column=1)
        self.samples_frame.grid(row=self.grid_size+2, column=0, columnspan=3)

        self.button_frame = tk.Frame(self)
        (HoverButton(
            self.button_frame, text="Confirm", width=10, command=self.confirm
                    ).pack(side='left', padx=5, pady=5)
        )
        (HoverButton(
            self.button_frame, text="Cancel", width=10,
            command=lambda: self.destroy()
                    ).pack(side='left', padx=5, pady=5)
        )
        self.button_frame.grid(row=self.grid_size+3, column=0, columnspan=3)

        self.title("Production")
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        # Updates raw_materials, production_log and bottles/samples or
        # in_progress sql tables with corresponding values from
        # production toplevel.
        self.product_amount = (self.product_frame
                               .grid_slaves(row=1, column=0)[0]
                               .get()
                              )
        self.case_amount = (self.product_frame
                            .grid_slaves(row=1, column=1)[0]
                            .get()
                           )
        self.product_var = (self.product_frame
                            .grid_slaves(row=1, column=2)[0]
                            .get()
                           )
        self.samples_var = self.samples_entry.get()
        # Raw material options.
        self.materials = [x.get() for x
                          in reversed(self.grid_slaves())
                          if (x.winfo_class() == 'TCombobox')]
        # Raw material entries.
        self.entries = [x.get() for x
                        in reversed(self.grid_slaves())
                        if (x.winfo_class() == 'Entry')]
        # Raw material product types.
        self.types = [x.cget("text").rstrip(":") for x
                      in reversed(self.grid_slaves())
                      if (x.winfo_class() == 'Label'
                      and x.cget("text").find(":") != -1)]
        for (entry, material) in zip(self.entries, self.materials):
            # Check material inputs to ensure non-empty values.
            if (not entry and material != "None"):
                messagebox.showerror(
                    "Materials Input Error",
                    "At least one input within the materials used section is " +
                    "blank, please try again.", parent=self)
                return
        # Check product and case amounts.
        if (not self.product_amount) or (not self.case_amount):
            messagebox.showerror(
                "Product Input Error",
                "One or more of the 'Total Bottles' or 'Cases' entries are " +
                "blank, please try again.", parent=self)
            return
        # Check sample amount and 'finished' checkbox.
        if (not self.samples_entry.get()) and (self.check_var.get() == 1):
            messagebox.showerror(
                "Sample Input Error",
                "The samples entry must be non-empty, please try again.",
                parent=self)
            return
        self.curr_date = date.today()
        # Update 'in_progress' table if products checked as unfinished.
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        if self.check_var.get() == 0:
            self.query = "INSERT INTO in_progress VALUES (%s,%s,%s,%s)"
            self.cur.execute(self.query, (self.curr_date, self.product_var,
                                          self.product_amount, self.desc_var))
        # Update 'bottles' and 'samples' tables if products are
        # considered finished
        elif self.check_var.get() == 1:
            self.query = ("UPDATE bottles SET amount=(amount + %s) " +
                          "WHERE product=%s")
            self.cur.execute(self.query, (self.case_amount, self.product_var))
            self.query = ("UPDATE samples SET amount=(amount + %s) " +
                          "WHERE product=%s")
            self.cur.execute(self.query, (self.samples_var, self.product_var))
        # Update 'production log'.
        self.query = "INSERT INTO production VALUES (%s,%s,%s)"
        self.cur.execute(self.query, (self.curr_date, self.product_var,
                                      self.product_amount))
        # Update 'raw_materials' data.
        for (material, subtr,type) in zip(self.materials, self.entries,
                                          self.types):
            if material != "None":
                self.query = ("UPDATE raw_materials " +
                              "SET amount=GREATEST((amount - %s), 0) " +
                              "WHERE product=%s AND type=%s")
                self.cur.execute(self.query, (subtr, material, type))
        self.conn.commit()
        self.cur.close()
        self.conn.close()
        db_update()
        view_products('raw_materials', 'All', 'All', self.gui_table)
        self.destroy()

    def cbox_check(self):
        # Activate sample entry box and cases amount entry if checkbox
        # is checked.
        if self.check_var.get() == 1:
            self.samples_entry.config(state='normal')
            self.samples_entry.delete(0, 'end')
            (self.product_frame.grid_slaves(row=1, column=1)[0]
             .config(state='normal'))
            self.product_frame.grid_slaves(row=1, column=1)[0].delete(0, 'end')
        # Disable sample entry box and cases amount entry if checkbox is
        # unchecked.
        if self.check_var.get() == 0:
            self.samples_entry.delete(0, 'end')
            self.samples_entry.insert(0, "0")
            self.samples_entry.config(state='readonly')
            self.product_frame.grid_slaves(row=1, column=1)[0].insert(0, "0")
            (self.product_frame.grid_slaves(row=1, column=1)[0]
             .config(state='readonly'))

            def desc_cancel():
                # Sets cbox value to 1 to prevent user from continuing
                # without entering description value or entering sample
                # amount.
                self.desc_tl.destroy()
                self.check_var.set(1)
                self.cbox_check()


            def desc_set():
                # Set description variable based on description text.
                # Check if textbox is empty
                if not self.desc_text.compare("end-1c", "==", "1.0"):
                    self.desc_var = self.desc_text.get("1.0", 'end')
                    self.desc_tl.destroy()
                else:
                    messagebox.showerror("Input Error",
                                         "Please input a description.",
                                         parent=self.desc_tl)
                    return

            # Toplevel to insert description for in_progress table.
            self.desc_var = tk.StringVar()
            self.desc_tl = tk.Toplevel(self)
            (tk.Message(
                self.desc_tl,
                text="Please provide a description of why " +
                     "the production was considered unfinished. " +
                     "(ex. 'bottles unlabeled', 'waiting for labels')",
                     width=300
                       ).grid(row=0, column=0, columnspan=2)
            )
            self.desc_text = tk.Text(self.desc_tl, height=2, width=30)
            self.desc_text.grid(row=1, column=0, columnspan=2)
            self.desc_fr = tk.Frame(self.desc_tl)
            self.conf_b = HoverButton(self.desc_fr, text="Confirm",
                                      command=desc_set)
            self.conf_b.grid(row=0, column=0)
            (HoverButton(self.desc_fr, text="Cancel", command=desc_cancel
                        ).grid(row=0, column=1)
            )
            self.desc_fr.grid(row=2, column=0, columnspan=2)

            # Prevent use of 'x-out' button.
            self.desc_tl.protocol("WM_DELETE_WINDOW", disable_event)
            self.desc_tl.title("Production Description")
            self.desc_tl.resizable(0,0)
            self.desc_tl.geometry("+%d+%d" % (self.x + 30, self.y + 30))
            self.desc_tl.focus()
            # Prevent user from clicking outside of toplevel.
            self.desc_tl.grab_set()


class PurchaseOrderView(tk.Toplevel):
    # Toplevel used to create purchase orders.  Updates purchase_order
    # and bottles/samples based on values retrieved from the toplevel.
    def __init__(self, master):
        self.master = master
        self.x = (screen_width/2) - 350
        self.y = (screen_height/2) - 350
        # Products used for product combobox
        self.conn = dbx_aws.db_conn()
        self.query = "SELECT product FROM bottles"
        self.product_rows = dbx_aws.db_exec(self.conn, self.query)
        self.product_rows = [x[0] for x in self.product_rows.data]
        self.product_rows.append("")
        tk.Toplevel.__init__(self, master=self.master)

        self.title_fr = tk.Frame(self)
        (tk.Label(self.title_fr, text="Purchase Order", font="Arial 10 bold"
                 ).pack()
        )
        self.title_fr.grid(row=0, column=0, columnspan=2)

        # Frame containing purchase order shipment information.
        self.info_fr = tk.Frame(self, pady=2)
        tk.Label(self.info_fr, text="From:").grid(row=0, column=0, sticky='W')
        tk.Label(self.info_fr, text="PO Number:").grid(row=1, column=0,
                                                       sticky='W')
        tk.Label(self.info_fr, text="To:").grid(row=2, column=0, sticky='W')
        tk.Label(self.info_fr, text="PO Date:").grid(row=0, column=2,
                                                     sticky='W')
        tk.Label(self.info_fr, text="Pick Up Date:").grid(row=1, column=2,
                                                          sticky='W')
        tk.Entry(self.info_fr, justify='center').grid(row=0, column=1)

        # Search for last purchase order in corresponding year folder
        # if none exists, create this year's folder.
        self.year = str(datetime.now().year)
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        self.query = (
            "SELECT po_number from purchase_orders " +
            "WHERE po_number LIKE \'{}%\'").format(self.year)
        self.cur.execute(self.query)
        self.completed_pos = [x[0] for x in self.cur.fetchall()]
        self.query = (
            "SELECT po_number from pending_po " +
            "WHERE po_number LIKE \'{}%\'").format(self.year)
        self.cur.execute(self.query)
        self.pending_pos = [x[0] for x in self.cur.fetchall()]
        self.cur.close()
        self.conn.close()
        self.all_pos = self.completed_pos + self.pending_pos
        # Update po-number entry with latest po number + 1.
        if self.all_pos:
            self.max_po = 0
            for po in self.all_pos:
                if int(po[5:]) > self.max_po:
                    self.max_po = int(po[5:])
            self.new_po_num = "{}-{:03}".format(self.year, self.max_po + 1)
        else:
            self.new_po_num = "{}-001".format(self.year)

        self.po_entry = tk.Entry(self.info_fr, justify='center')
        self.po_entry.insert(0, self.new_po_num)
        self.po_entry.grid(row=1, column=1)
        tk.Entry(self.info_fr, justify='center').grid(row=2, column=1)
        self.po_date = tk.Entry(self.info_fr, justify='center',
                                state='readonly')
        self.po_date.grid(row=0, column=3)
        self.cal_photo = ImageTk.PhotoImage(cal_image)
        self.po_cal_link = HoverButton(
            self.info_fr, image=self.cal_photo,
            command=lambda: cal_button(self, self.po_date, self.master)
        )
        self.po_cal_link.image = self.cal_photo
        self.po_cal_link.grid(row=0,column=4)
        self.pu_date = tk.Entry(self.info_fr, justify='center',
                                state='readonly')
        self.pu_date.grid(row=1, column=3)
        self.pu_cal_link = HoverButton(
            self.info_fr, image=self.cal_photo,
            command=lambda: cal_button(self, self.pu_date, self.master)
        )
        self.pu_cal_link.image = self.cal_photo
        self.pu_cal_link.grid(row=1,column=4)
        self.info_fr.grid(row=1, column=0, columnspan=2)

        # Frame containing purchase order product information.
        self.order_fr = tk.Frame(self, padx=33)
        tk.Label(self.order_fr, text="QTY").grid(row=0, column=0,
                                                 sticky='NESW')
        tk.Label(self.order_fr, text="UNIT").grid(row=0, column=1,
                                                  sticky='NESW')
        tk.Label(self.order_fr, text="PRODUCT").grid(row=0, column=2,
                                                     sticky='NESW')
        tk.Label(self.order_fr, text="UNIT COST").grid(row=0, column=3,
                                                       sticky='NESW')
        tk.Label(self.order_fr, text="TOTAL").grid(row=0, column=4,
                                                   sticky='NESW')
        for i in range(1,19): # Purchase order product information.
            (tk.Entry(
                self.order_fr, width=5, justify="center", validate="key",
                validatecommand=(self.register(valid_dig),'%S','%d')
                     ).grid(row=i, column=0, sticky='NESW')
            )
            (ttk.Combobox(
                self.order_fr, values=['Cases','Bottles', ''],
                width=7, justify="center", state='readonly'
                         ).grid(row=i, column=1, sticky='NESW')
            )
            (ttk.Combobox(
                self.order_fr, values=self.product_rows, justify="center",
                state='readonly'
                         ).grid(row=i, column=2)
            )
            (tk.Entry(
                self.order_fr, width=12, justify="center", validate="key",
                validatecommand=(self.register(valid_dec),'%S','%s','%d')
                     ).grid(row=i, column=3, sticky='NESW')
            )
            (tk.Entry(
                self.order_fr, width=12, justify="center", bg="light gray"
                     ).grid(row=i,column=4,sticky='NESW')
            )
        (tk.Label(
            self.order_fr, text="TOTAL", background="dark slate gray",
            relief="raised", fg="white"
                 ).grid(row=19, column=0, columnspan=5, sticky='EW')
        )
        self.total_var = tk.StringVar()
        self.total_label = tk.Label(
            self.order_fr, background="gray", fg="white",
            width=10, relief="raised", textvariable=self.total_var
        )
        self.total_label.grid(row=19, column=4, sticky='EW')
        for label in self.order_fr.grid_slaves(row=0):
            label.config(background="dark slate gray", relief="raised",
                         fg="white")
        self.order_fr.grid(row=2, column=0, columnspan=2, pady=2)

        self.check_var = tk.IntVar()
        self.check_var.set(1)
        self.check_b = tk.Checkbutton(
            self, text="Pending Purchase Order", variable=self.check_var
        )
        self.check_b.grid(row=3, column=0, columnspan=2)

        self.btn_fr = tk.Frame(self)
        (HoverButton(
            self.btn_fr, text="Confirm", command=self.confirm
                    ).grid(row=0, column=0, padx=10)
        )
        (HoverButton(
            self.btn_fr, text="Cancel", command=lambda: self.destroy()
                    ).grid(row=0, column=1, padx=10)
        )
        self.btn_fr.grid(row=4, column=0, columnspan=2, pady=2)

        self.total_after()
        self.geometry("%dx%d+%d+%d" % (464, 625, self.x, self.y))
        self.focus()

    def total_after(self):
        # Updates total column entry values to be product of quantity
        # and price. Sums total columns into final total column,
        # total_var
        self.total_entries = [x for x
                              in reversed(self.order_fr.grid_slaves(column=4))
                              if x.winfo_class() == 'Entry']
        self.total_sum = 0
        for (entry, i) in zip(self.total_entries, range(1, 19)):
            entry.delete(0, 'end')
            try:
                # (quantity * unit cost)
                self.row_total = round(
                  Decimal(self.order_fr.grid_slaves(row=i, column=0)[0].get())
                  *Decimal(self.order_fr.grid_slaves(row=i, column=3)[0].get()),
                  2)
                entry.insert(0, self.row_total)
                self.total_sum += self.row_total
            except:
                pass
        self.total_sum = round(self.total_sum, 2)
        self.total_var.set(self.total_sum)
        self.after_func = self.after(150, self.total_after)

    def confirm(self):
        self.open_ques = messagebox.askquestion(
            "Purchase Order Confirmation",
            "Are you sure you want to confirm? Please make sure everything is "+
            "entered correctly. Confirming will update inventory values and " +
            "save the purchase order with the file name, " +
            self.new_po_num + ".xlsx. \n \n" +
            "If the 'Pending Purchase Order' checkbox is checked, " +
            "the purchase order will be stored within the 'Pending' tab to " +
            "be completed or removed at a later date.", parent=self)

        if self.open_ques == 'no':
            return self.total_after()
        else:
            # Stop after_func to ensure total_var value.
            self.after_cancel(self.after_func)
            self.info_entries = [x.get() for x
                                 in reversed(self.info_fr.grid_slaves())
                                 if x.winfo_class() == 'Entry']
            for entry in self.info_entries:
                if not entry:
                    messagebox.showerror(
                        "Input Error",
                        "Please make sure all of the top entries have values.",
                        parent=self)
                    return self.total_after()
            self.po_entries = [x.get() for x
                               in reversed(self.order_fr.grid_slaves())
                               if (x.winfo_class() == 'Entry'
                               or x.winfo_class() == "TCombobox")]
            # List of lists containing po order values.
            self.complete_po_lists = [self.po_entries[x:x+5] for x
                                      in range(0, len(self.po_entries), 5)]
            for list in self.complete_po_lists:
                if any(list) and not all(list):
                    messagebox.showerror(
                        "Input Error",
                        "Please make sure all of the purchase order entries " +
                        "are fully complete.", parent=self)
                    return self.total_after()
                else:
                    continue

            if self.check_var.get() == 1:
                # Add purchase orders to 'pending_po' table.
                self.conn = dbx_aws.db_conn()
                self.cur = self.conn.cursor()
                for po_list in (x for x in self.complete_po_lists if all(x)):
                    self.query = ("INSERT INTO pending_po " +
                                  "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)")
                    self.cur.execute(
                        self.query,
                        (self.info_entries[3], self.info_entries[4], po_list[2],
                         po_list[0], po_list[1], po_list[3], po_list[4],
                         self.info_entries[2], self.po_entry.get()))
                self.conn.commit()
                self.cur.close()
                self.conn.close()
            else:
                confirm_po(self.info_entries, self.complete_po_lists,
                           self.po_entry.get())
            self.destroy()


class FinishView(tk.Toplevel):
    # Toplevel used to finish 'in progress' production values.  Deletes
    # in_progress table entry and updates bottles/samples upon
    # confirmation.
    def __init__(self, master, values):
        self.master = master
        self.values = values
        self.x = x + 150
        self.y = y + 20
        self.product = self.values[1]
        tk.Toplevel.__init__(self, master=self.master)

        self.title_fr = tk.Frame(self)
        (tk.Label(
            self.title_fr, text="Finish Production", font="Arial 10 bold",
            pady=5
                 ).pack()
        )
        self.title_fr.grid(row=0, column=0, columnspan=2, pady=5)

        self.prod_fr = tk.Frame(self)
        (tk.Label(
            self.prod_fr, text=self.product
                 ).grid(row=0, column=0, columnspan=2)
        )
        (tk.Label(
            self.prod_fr, text="Cases:"
                 ).grid(row=1, column=0, sticky=W, pady=2)
        )
        (tk.Entry(
            self.prod_fr, validate="key",
            validatecommand=(self.register(valid_dig),'%S','%d')
                 ).grid(row=1, column=1, pady=2)
        )
        (tk.Label(
            self.prod_fr, text="Samples:"
                 ).grid(row=2, column=0, sticky=W, pady=2)
        )
        (tk.Entry(
            self.prod_fr, validate="key",
            validatecommand=(self.register(valid_dig),'%S','%d')
                 ).grid(row=2, column=1, pady=2)
        )
        self.prod_fr.grid(row=1, column=0, columnspan=2)

        self.button_fr = tk.Frame(self)
        (HoverButton(
            self.button_fr, text="Confirm", command=self.confirm
                    ).pack(side='left')
        )
        (HoverButton(
            self.button_fr, text="Cancel", command=lambda: self.destroy()
                    ).pack(side='left')
        )
        self.button_fr.grid(row=2, column=0, columnspan=2)

        self.title("Finish {} Production".format(self.product))
        self.geometry("%dx%d+%d+%d" % (178, 140, self.x, self.y))
        self.resizable(0,0)
        self.focus()

    def confirm(self):
        # Update 'bottles' and 'samples' tables with newly created
        # products. Remove info from 'in_progress' table.
        self.conf_quest = messagebox.askquestion(
            "Finish this product?",
            "Are you sure you want to confirm? Make sure everything is " +
            "entered correctly before continuing.", parent=self)
        if self.conf_quest == "yes":
            self.entries = [x.get() for x
                            in reversed(self.prod_fr.grid_slaves())
                            if x.winfo_class() == "Entry"]
            if all(self.entries):
                self.conn = dbx_aws.db_conn()
                self.cur = self.conn.cursor()
                self.query = ("UPDATE bottles SET amount=(amount + %s) " +
                              "WHERE product=%s")
                cur.execute(self.query, (self.entries[0], self.product))
                self.query = ("UPDATE samples SET amount=(amount + %s) " +
                              "WHERE product=%s")
                cur.execute(self.query, (self.entries[1], self.product))
                self.query = "DELETE FROM in_progress WHERE product=?"
                cur.execute(self.query, (self.product,))
                self.conn.commit()
                self.cur.close()
                self.conn.close()

                db_update()
                view_products('in_progress', 'All', 'All', inprog_tbl)
                self.destroy()
            else:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the info entries have values.",
                    parent=self)
                return
        else:
            return


class MashProductionView(tk.Toplevel):
    # Toplevel to input values for mash production.
    def __init__(self, master, mash_table):
        self.master = master
        self.mash_table = mash_table
        self.x = x + 150
        self.y = y + 150
        # Populated with lists of length 3 (grain, amt, order #) in grain_recur
        self.grain_info_tbl = []
        tk.Toplevel.__init__(self, master=self.master)
        # Get previous mash information.

        try:
            self.conn = dbx_aws.db_conn()
            self.query = ("SELECT mash_no, type FROM mashes " +
                          "ORDER BY mash_no DESC LIMIT 1")
            self.mash_data = dbx_aws.db_exec(self.conn, self.query)
            self.prev_mash = self.mash_data.data[0][0]
            self.prev_mash_num = self.prev_mash[0]
            self.prev_mash_type = self.prev_mash[1]
            # Mash number regex matches.
            self.mo = mashRegex.search(self.prev_mash_num)
            self.year = self.mo.group(1) # Prev mash's year.
            self.mash_count = self.mo.group(5) # Prev mash's ID number.
            self.mash_letter = self.mo.group(6) # Prev mash's letter variable.
            self.mash_letters = list(string.ascii_uppercase[:8]) # Letters A-H
        except:
            self.year = int(datetime.now().year)
            self.mash_count = 0
            self.mash_letter = "A"
            self.prev_mash_type = None
        self.conn = dbx_aws.db_conn()
        self.query = "SELECT type, order_number FROM grain"
        self.grain_ord_nums = dbx_aws.db_exec(self.conn, self.query).data
        self.grain_ord_dict = {}
        for (key, value) in self.grain_ord_nums:
            if key in self.grain_ord_dict:
                self.grain_ord_dict[key].append(value)
            else:
                self.grain_ord_dict[key] = [value]

        self.title_fr = tk.Frame(self)
        (tk.Label(
            self.title_fr, text="Mash Production", font="Arial 10 bold", pady=5
                 ).pack()
        )
        self.title_fr.grid(row=0, column=0, columnspan=3)

        self.type_fr = tk.Frame(self)
        tk.Label(self.type_fr, text="Mash Type:").grid(row=0, column=0)
        self.type_menu = ttk.Combobox(
            self.type_fr, values=["Bourbon","Rye","Malt","Rum"],
            width=16, justify="center", state="readonly")
        self.type_menu.set("Bourbon")
        self.type_menu.bind("<<ComboboxSelected>>", self.tplvl_upd)
        self.type_menu.grid(row=0, column=1)
        tk.Label(self.type_fr, text="Mash Number:").grid(row=1, column=0)
        self.mash_num_entry = tk.Entry(self.type_fr, justify='center')
        self.mash_num_entry.grid(row=1, column=1)
        tk.Label(self.type_fr, text="Date:").grid(row=2, column=0)
        self.date = tk.StringVar()
        # Update mash number with month value.
        self.date.trace(
            "w",
            lambda name, index, mode:
            self.mash_num_upd(self.prev_mash_type, self.type_menu.get()))
        self.date_entry = tk.Entry(
            self.type_fr, state="readonly", justify="center",
            textvariable=self.date)
        self.date_entry.grid(row=2,column=1)
        self.cal_photo = ImageTk.PhotoImage(cal_image)
        self.cal_link = HoverButton(
            self.type_fr, image=self.cal_photo,
            command=lambda:
            cal_button(self, self.date_entry, self.master)
        )
        self.cal_link.image = self.cal_photo
        self.cal_link.grid(row=2, column=2)
        self.type_fr.grid(row=1, column=0, columnspan=3)

        self.grain_fr = tk.Frame(self, pady=5, padx=5, height=100, width=340)
        self.grain_fr.grid_propagate(0) # Frame size doesn't change

        self.button_fr = tk.Frame(self, padx=10)
        (HoverButton(
            self.button_fr, text="Confirm", command=self.confirm
                    ).grid(row=0, column=0)
        )
        (HoverButton(
            self.button_fr, text="Cancel", command=lambda: self.destroy()
                    ).grid(row=0, column=1)
        )
        self.button_fr.grid(row=3, column=0, columnspan=3)

        self.title("Mash Production")
        self.geometry("%dx%d+%d+%d" % (350, 240, self.x, self.y))
        self.resizable(0,0)
        self.focus()
        self.type_menu.event_generate("<<ComboboxSelected>>")

    def mash_num_upd(self, prev_type, curr_type):
        # Update mash number entry based on current grain type and
        # previous mash type.
        self.mash_num_entry.delete(0, 'end')
        if self.date_entry.get():
            self.month = self.date_entry.get()[5:7]
        else:
            self.month = '{:02d}'.format(datetime.now().month)
        # ex 2019/03-4A
        self.new_batch_num = (str(self.year) + "/" + str(self.month) + "-"
                              + str(int(self.mash_count) + 1) + "A")
        # Handle new year case.
        if int(self.year) != int(datetime.now().year):
            self.mash_num_entry.insert(0, self.year + "/1-1A")
        else:
            if prev_type == curr_type:
                if self.mash_letter != "H": # Same type, same batch case.
                    self.mash_let_indx = (self.mash_letters
                                          .index(self.mash_letter) + 1)
                    self.new_batch_num = (
                        str(self.year) + "/" + str(self.month)
                        + "-" + str(self.mash_count)
                        + self.mash_letters[self.mash_let_indx])
                    self.mash_num_entry.insert(0, self.new_batch_num)
                else:   # Same type, next batch case.
                    self.mash_num_entry.insert(0, self.new_batch_num)
            else:   # New type, new batch case.
                self.mash_num_entry.insert(0, self.new_batch_num)

    def fill_frame(self, gr_lst):
        # Update grain frame with grain types used for selected mash
        # type.
        for index, grain in enumerate(gr_lst,1):
            tk.Label(self.grain_fr, text=grain).grid(row=index, column=0)
            (tk.Entry(
                self.grain_fr, validate='key',
                validatecommand=(self.register(valid_dig), '%S', '%d')
                     ).grid(row=index, column=1)
            )
            (ttk.Combobox(
                self.grain_fr, values=self.grain_ord_dict[grain], width=16,
                justify='center', state='readonly'
                         ).grid(row=index, column=2, padx=3)
            )

    def tplvl_upd(self, event):
        # Remove grain inputs and replace with new ones corresponding to
        # grain type.  Update mash number based on previous and current
        # mash types.
        self.type = self.type_menu.get()
        self.mash_num_upd(self.prev_mash_type, self.type)
        for widg in self.grain_fr.grid_slaves():
            widg.grid_forget()
        (tk.Label(
            self.grain_fr, text="Grain", font="Arial 10 bold"
                 ).grid(row=0, column=0, columnspan=3)
        )
        try:
            if self.type == "Bourbon":
                self.fill_frame(["Corn","Rye","Malted Barley"])
            elif self.type == "Rye":
                self.fill_frame(["Rye","Malted Wheat"])
            elif self.type == "Malt":
                self.fill_frame(["Malted Barley","Wheat","Oat"])
            else:
                self.fill_frame(["Molasses"])
        except KeyError:
            self.destroy()
            messagebox.showerror(
                "Grain Inventory Error",
                "There was an issue retrieving the purchase order number for " +
                "one of the grain types required for this type of mash. " +
                "Please make sure there is an inventory value for each type " +
                "of grain needed in this mash. \n \n" +
                "If you are receiving this message upon opening " +
                "'Produce Mash', there is a missing grain for Bourbon.")
            return
        self.grain_fr.grid(row=2, column=0, columnspan=3)

    def confirm(self):
        # Subtract grain used from 'grain' table, update 'mash_log'
        # based on entries. Create a 'mash log' word file with certain
        # info filled out.
        self.grain_types = [x.cget("text") for x
                            in reversed(self.grain_fr.grid_slaves())
                            if x.winfo_class() == "Label"][1:]
        self.grain_amts = [x.get() for x
                           in reversed(self.grain_fr.grid_slaves())
                           if x.winfo_class() == "Entry"]
        self.type_entries = [x.get() for x
                             in reversed(self.type_fr.grid_slaves())
                             if x.winfo_class() == "Entry"]
        self.order_nums = [x.get() for x
                           in reversed(self.grain_fr.grid_slaves())
                           if x.winfo_class() == "TCombobox"]
        self.entry_check = self.grain_amts + self.type_entries + self.order_nums
        for entry in self.entry_check:
            if not entry:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the entries are fully " +
                    "complete.", parent=self)
                return
        try:
            self.file_path = "{}/production_sheets/Mash_Log.docx".format(
                                                                    os.getcwd())
            self.file = open(self.file_path, 'rb')
            self.document = Document(self.file)
            self.file.close()
        except:
            messagebox.showerror(
                "File Error",
                "It seems the Word Document you are trying to edit or change " +
                "is already open, please close it and try again.",
                parent=self)
            return

        # Subtract grain amounts from inventory.
        for (type, amount, order_num) in zip(self.grain_types, self.grain_amts,
                                             self.order_nums):
            self.grain_recur(type, amount, order_num)
        self.conn = dbx_aws.db_conn()
        self.query = "INSERT INTO mashes VALUES (%s,%s,%s,%s)"
        dbx_aws.db_exec(self.conn, self.query, False, self.date_entry.get(),
                        self.mash_num_entry.get(), self.type_menu.get(),
                        ", ".join(self.order_nums))

        # Word doc tables
        self.info_table = self.document.tables[0]
        self.grain_table = self.document.tables[1]
        # Write info for top of file
        for (row, info) in zip(self.info_table.rows,
                              [self.date_entry.get(),
                               self.type_menu.get(),
                               self.mash_num_entry.get()]):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = info
                    for run in para.runs:
                        run.font.name = "Verdana"
                        run.font.size = Pt(14)
        # Write info for grain table section of file
        for (row, gr_list) in zip(self.grain_table.rows, self.grain_info_tbl):
            for (cell, num) in zip(row.cells, range(3)):
                for para in cell.paragraphs:
                    para.text = gr_list[num]
                    for run in para.runs:
                        run.font.name = "Verdana"
                        run.font.size = Pt(14)

        self.temp_file = "{}/temp_files/Last_Mash_Log.docx".format(os.getcwd())
        self.document.save(self.temp_file)

        self.open_ques = messagebox.askquestion(
            "Open the Mash Word File?",
            "Would you like to open the Mash Log file in Word? This will " +
            "allow you to print it now.")
        if self.open_ques == "yes":
            try:
                os.system('start ' + self.temp_file)
            except:
                messagebox.showerror(
                    "Program Error",
                    "There was an error opening Word.", parent=self)
        else:
            pass

        db_update()
        view_products('grain', 'All', 'All', grain_tbl)
        self.destroy()

    def grain_recur(self, type, amount, order_num, first=True):
        # Subtract amounts from respective grain. Recursion occurs when
        # a grain amount is zeroed out. User is asked for a second order
        # number for that type of grain to subtract from.
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        self.query = ("SELECT amount, date FROM grain " +
                      "WHERE type=%s AND order_number=%s")
        self.cur.execute(self.query, (type, order_num))
        self.grain_lst = list(self.cur.fetchall())
        self.grain_amt = (self.grain_lst[0][0])
        self.grain_date = (self.grain_lst[0][1])
        if self.grain_amt:
            self.grain_diff = int(self.grain_amt) - int(amount)
            if self.grain_diff > 0:
                self.query = ("UPDATE grain SET amount=%s, total=(%s * price) "+
                              "WHERE type=%s AND order_number=%s")
                self.cur.execute(self.query, (self.grain_diff, self.grain_diff,
                                              type, order_num))
            else:
                self.query = ("DELETE FROM grain WHERE type=%s " +
                              "AND order_number=%s")
                self.cur.execute(self.query, (type, order_num))
                self.query = "INSERT INTO grain_log VALUES (%s,%s,%s,%s)"
                self.cur.execute(self.query, (self.grain_date,
                                              self.date_entry.get(), type,
                                              order_num))
            self.conn.commit()
            self.cur.close()
            self.conn.close()
            if first == True:
                self.grain_info_tbl.append([type, str(amount), str(order_num)])
            elif first == False:
                for lst in self.grain_info_tbl:
                    if lst[0] == type:
                        lst[2] = lst[2] + ", " + order_num
                        self.order_nums.append(order_num)
            if self.grain_diff < 0:
                self.grain_diff = abs(self.grain_diff)
                self.grain_recur_tplvl(type, order_num)
        else:
            messagebox.showerror(
                "Grain Error",
                "There doesn't seem to be an inventory value for " + type +
                ", or there isn't enough grain, please fix this and try again.",
                parent=self)
            raise ValueError("Grain Error")

    def grain_recur_tplvl(self, type, order_num):
        # Toplevel to add another order number to fulfill grain amount
        self.combox_values = [x for x
                              in self.grain_ord_dict[type]
                              if x != order_num]
        if self.combox_values:
            self.recur_tplvl = tk.Toplevel(self)
            (tk.Label(
                self.recur_tplvl,
                text = ("There is {} lbs. of {} left to be used, \n " +
                        "please select another order number to subtract it "+
                        "from").format(self.grain_diff, type)
                       ).grid(row=0, column=0, columnspan=2, pady=5)
            )
            self.new_ord_box = ttk.Combobox(self.recur_tplvl,
                                            values=self.combox_values)
            self.new_ord_box.config(width=16, background="white",
                                      justify='center', state='readonly')
            self.new_ord_box.set(self.combox_values[0])
            self.new_ord_box.grid(row=1, column=0, columnspan=2, pady=5)
            (HoverButton(
                self.recur_tplvl, text="Confirm",
                command=lambda: self.recur_confirm(type)
                        ).grid(row=2, column=0, pady=5)
            )
            (HoverButton(
                self.recur_tplvl, text="Cancel",
                command=lambda: self.recur_cancel()
                        ).grid(row=2, column=1, pady=5)
            )

            self.recur_tplvl.protocol("WM_DELETE_WINDOW", disable_event)
            self.recur_tplvl.title("Production Description")
            self.recur_tplvl.resizable(0,0)
            self.recur_tplvl.geometry("+%d+%d" % (self.x + 60, self.y + 60))
            self.recur_tplvl.focus()
            self.recur_tplvl.grab_set()
            self.wait_window(self.recur_tplvl)
        else:
            messagebox.showerror(
                "Grain Error",
                ("There doesn't seem to be an inventory value for {}, or " +
                 "there isn't enough grain, please fix this and try again."
                ).format(type), parent=self)
            raise ValueError("Grain Error")

    def recur_cancel(self):
        # Button function to destroy toplevel and close sql connection
        self.recur_tplvl.destroy()
        raise ValueError("Grain Error")

    def recur_confirm(self, type):
        # Button function to destroy toplevel and run grain_recur with
        # new inputs
        self.grain_recur(type, self.grain_diff, self.new_ord_box.get(),
                         first=False)
        self.recur_tplvl.destroy()

class ReportsFrame(tk.Frame):
    # Creates frame containing inventory information from each sql
    # table.
    def __init__(self, master):
        self.master = master
        self.cur_year = datetime.now().year
        self.cur_month_ind = datetime.now().month
        self.year_choices = list(range(2019, self.cur_year + 1))
        self.month_choices = list(calendar.month_abbr)
        tk.Frame.__init__(self, master)

        self.year_fr = tk.LabelFrame(self, text="YEAR", font="Arial 8 bold")
        self.year_cmbo_box = ttk.Combobox(
            self.year_fr, values=self.year_choices, state='readonly',
            justify='center'
        )
        self.year_cmbo_box.set(self.cur_year)
        self.year_cmbo_box.bind("<<ComboboxSelected>>", self.year_upd)
        self.year_cmbo_box.pack(padx=5, pady=5)
        self.year_fr.grid(row=0, column=0, padx=5, pady=5)

        self.month_fr = tk.LabelFrame(self, text="MONTH", font="Arial 8 bold")
        self.month_cmbo_box = ttk.Combobox(
            self.month_fr, values=self.month_choices, state='readonly',
            justify='center'
        )
        self.month_cmbo_box.set(self.month_choices[self.cur_month_ind])
        self.month_cmbo_box.bind("<<ComboboxSelected>>", self.month_upd)
        self.month_cmbo_box.pack(padx=5, pady=5)
        self.month_fr.grid(row=0, column=1, padx=5, pady=5)

        self.invent_fr = tk.LabelFrame(
            self, text="Inventory", font="Arial 12 bold", fg="dark slate gray"
        )
        self.barrel_fr = tk.LabelFrame(
            self, text="Barrels", font="Arial 12 bold", fg="dark slate gray"
        )
        self.po_fr = tk.LabelFrame(
            self, text="Purchase Orders", font="Arial 12 bold",
            fg="dark slate gray"
        )
        self.logo_fr = tk.LabelFrame(
            self, font="Arial 10 bold", fg="dark slate gray"
        )
        self.logo_path = "ADCO_Logo.jpg"
        self.img = Image.open(self.logo_path)
        self.img = self.img.resize((640,520))
        self.img = ImageTk.PhotoImage(self.img)
        self.logo_label = tk.Label(self.logo_fr, image=self.img)
        self.logo_label.image = self.img
        self.logo_label.grid(row=0, column=0)
        self.logo_fr.grid(row=0, column=2, rowspan=5, sticky='nesw', pady=3,
                          padx=10)

        self.year_cmbo_box.event_generate("<<ComboboxSelected>>")

    def year_upd(self, event):
        # Generate combobox selection event for months, which will make
        # changes to values displayed.
        monthly_reports_update()
        self.year_sel = self.year_cmbo_box.get()
        self.month_cmbo_box.event_generate("<<ComboboxSelected>>")

    def month_upd(self, event):
        # Retrieve and display values from 'monthly_reports' table based
        # on year and month selected.
        self.month_sel = self.month_cmbo_box.get()
        self.month_sel = self.month_choices.index(self.month_sel)
        self.month_sel = "{:02}".format(self.month_sel)
        self.date_sel = "{}-{}-01".format(self.year_sel, self.month_sel)
        for widg in [*self.invent_fr.grid_slaves(), # * unpacks contents
                     *self.barrel_fr.grid_slaves(), # of iterables
                     *self.po_fr.grid_slaves()]:
            widg.grid_forget()
        # Inventory Values
        self.query = "SELECT * FROM monthly_reports WHERE date=%s"
        self.conn = dbx_aws.db_conn()
        self.monthly_totals = dbx_aws.db_exec(self.conn, self.query,
                                              True, self.date_sel)
        self.monthly_totals = [list(x) for x in self.monthly_totals.data]
        if self.monthly_totals:
            self.invent_vals_positions = {'raw_materials':0,
                                          'pending_cogs':1,
                                          'bottles':2,
                                          'samples':3,
                                          'pending_po':4,
                                          'grain':5}
            self.invent_vals = [x for x
                                in self.monthly_totals
                                if x[1] in self.invent_vals_positions.keys()]
            self.invent_vals.sort(key=lambda x:
                                  self.invent_vals_positions[x[1]])
            self.barrel_vals_positions = {'barreled_rum':0,
                                          'barreled_whiskey':1,
                                          'barrels':2}
            self.barrel_vals = [x for x
                                in self.monthly_totals
                                if x[1] in self.barrel_vals_positions.keys()]
            self.barrel_vals.sort(key=lambda x:
                                  self.barrel_vals_positions[x[1]])
            self.po_vals_positions = {'purchase_order_sales':0,
                                      'purchase_order_cogs':1,
                                      'pending_sales':2,
                                      'pending_cogs':3}
            self.po_vals = [x for x
                            in self.monthly_totals
                            if x[1] in self.po_vals_positions.keys()]
            self.po_vals.sort(key=lambda x: self.po_vals_positions[x[1]])
            self.monthly_frames_fill(self.invent_vals, self.invent_fr)
            self.monthly_frames_fill(self.barrel_vals, self.barrel_fr)
            self.monthly_frames_fill(self.po_vals, self.po_fr)
        else:
            (tk.Label(
                self.invent_fr, text="N/A", font="Arial 30 bold", fg="gray"
                     ).grid(row=0, column=0, columnspan=2)
            )
            (tk.Label(
                self.barrel_fr, text="N/A", font="Arial 30 bold", fg="gray"
                     ).grid(row=0, column=0, columnspan=2)
            )
            (tk.Label(
                self.po_fr, text="N/A", font="Arial 30 bold", fg="gray"
                     ).grid(row=0, column=0, columnspan=2)
            )
        self.invent_fr.grid(row=1, column=0, columnspan=2, padx=5, pady=5,
                           sticky="NESW")
        self.barrel_fr.grid(row=2, column=0, columnspan=2, padx=5, pady=5,
                            sticky="NESW")
        self.po_fr.grid(row=3, column=0, columnspan=2, padx=5, pady=5,
                            sticky="NESW")

    def monthly_frames_fill(self, inv_vals, inv_fr):
        self.invent_sum = 0
        self.grid_ind = 0
        self.neg_vals = ['purchase_order_cogs', 'pending_cogs']
        for index, lst in enumerate(inv_vals):
            self.invent_sum += float(lst[2])
            self.grid_ind += 1
            # ex. Purchase Orders:
            self.txt1 = lst[1].replace("_", " ").upper()
            self.txt1 = "{}:".format(self.txt1)
            # ex -29000 -> $29,000.00
            self.txt2 = "${0:,.2f}".format(float(lst[2])).replace("-","")
            (tk.Label(
                inv_fr, text=self.txt1, font="Arial 10 bold"
                     ).grid(row=index, column=0, padx=20, sticky="W")
            )
            self.total_label = tk.Label(
                inv_fr, text=self.txt2, font="Arial 10 bold", borderwidth=1,
                relief="solid", width=9
            )
            if float(lst[2]) < 0:
                self.total_label.config(bg='pink', fg='red')
            self.total_label.grid(row=index, column=1, ipadx=20, sticky="E")
        (tk.Label(
            inv_fr, text="TOTAL:", font="Arial 12 bold"
                 ).grid(row=self.grid_ind, column=0, padx=20, sticky="W")
        )
        (tk.Label(
            inv_fr, text="${0:,.2f}".format(self.invent_sum),
            font="Arial 12 bold", borderwidth=2, relief="solid", width=10
                 ).grid(row=self.grid_ind, column=1, ipadx=20, sticky="E")
        )
        # Force widget to fill column
        inv_fr.columnconfigure(0, weight=1)
        inv_fr.columnconfigure(1, weight=1)


class SheetLabel(tk.Label):
    # Creates a clickable label with link to file in given file location.
    def __init__(self, master, text, file_location):

        tk.Label.__init__(self, master, text=text, cursor="hand2",
                          font="Times 14 underline", fg="#0000EE")
        def button_click(event):
            # Changes label color from 'blue' to 'purple' and opens the
            # file.
            if self['fg'] =="#0000EE":
                self['fg'] = "#551A8B"
            else:
                self['fg'] = "#551A8B"
            file = webbrowser.open_new(file_location)
        self.bind("<Button-1>", func=button_click)


class CommandFrame(tk.Frame):
    # Creates frame on the left side of the treeview tables.  Used to
    # place command buttons for interacting with data.
    def __init__(self,master):
        self.master = master
        self.height = height
        self.width = command_width
        tk.Frame.__init__(self, master=self.master, height=self.height,
                       width=self.width)


class ViewFrame(tk.LabelFrame):
    # Frame that contains comboboxes for selecting filters within the
    # current displayed inventory table. Filters will cause the table
    # to display updated values based on the chosen filter.
    def __init__(self, master, sql_table, gui_table):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.height = int(screen_height/1.5)
        self.text_var = tk.StringVar()
        tk.LabelFrame.__init__(self, master=self.master, height=self.height,
                               bd=5, relief='ridge', text="View", font="bold")

        # Columns selector.
        tk.Label(self, text="Columns").grid(row=0, column=0)
        self.unused_filters = ["Price", "Total", "Amount", "Unit", "Case Size",
                               "Grains", "Mash No", "PO No.", "Proof Gallons",
                               "Order No"]
        self.column_vals = [x for x
                            in list(self.gui_table["columns"])
                            if (x not in self.unused_filters)
                           ]
        self.column_vals.append("All")
        self.columns = ttk.Combobox(self, values=self.column_vals)
        self.columns.set(self.column_vals[-1])
        self.columns.config(width=18, background="white", justify='center',
                            state='readonly')
        self.columns.bind("<<ComboboxSelected>>", self.col_upd)
        self.columns.grid(row=1, column=0, pady=3, padx=5)

        # Item selector.
        tk.Label(self, text="Values").grid(row=2, column=0)
        self.rows = ttk.Combobox(self, values=["N/A"])
        self.rows.set("N/A")
        self.rows.config(width=18, background="white", justify='center',
                         state='readonly')
        self.rows.bind("<<ComboboxSelected>>", self.row_upd)
        self.rows.grid(row=3, column=0, pady=3, padx=5)

        # Total value label.
        if any(s in self.gui_table.columns for s in ('Price', 'Proof Gallons')):
            if self.sql_table in ["purchase_orders", "pending_po"]:
                self.inv_label = "Transactions Value"
            else:
                self.inv_label = "Inventory Value"
            tk.Label(self, text=self.inv_label).grid(row=4, column=0)
            (tk.Label(
                self, textvariable=self.text_var, bd=10, font="Arial 15 bold",
                fg="dark slate grey"
                     ).grid(row=5, column=0)
            )
        self.pack()

    def col_upd(self, event):
        # Update row-combobox values based on column selection.
        self.column_val = self.columns.get().lower().replace(" ","_")
        if self.column_val == "all":
            self.value_rows = ["N/A"]
        else:
            self.conn = dbx_aws.db_conn()
            self.query = "SELECT {} FROM {}".format(self.column_val,
                                                    self.sql_table)
            self.data = dbx_aws.db_exec(self.conn, self.query).data

            if self.column_val == "barrel_no":
                self.value_rows = {"{}-XXX".format(x[0][:2]) for x
                                   in self.data}
            elif "date" in self.column_val: # Get year from date value
                    self.value_rows = {x[0].strftime("%Y") for x in self.data}
            elif self.column_val == "age":
                self.value_rows = {"{} year(s)".format(int(x[0] / 365)) for x
                                   in self.data}
            elif self.column_val == "product":
                self.value_rows = [x[0] for x in self.data]
                for (ind, item) in enumerate(self.value_rows):
                    self.mo = re.search("\d", item)
                    # Strip value at location of first digit.
                    if self.mo and self.mo.start() > 1:
                        self.value_rows[ind] = (
                            self.value_rows[ind][0:self.mo.start() - 1])
                    else:
                        continue
                self.value_rows = set(self.value_rows)
            else:
                try:
                    self.value_rows = {x[0] for x in self.data}
                    self.value_rows = set(self.value_rows)
                except IndexError:
                    self.value_rows = ["N/A"]
        try:
            self.value_rows = list(self.value_rows)
            self.value_rows.sort()
            self.rows.config(values=self.value_rows)
            self.rows.set(self.value_rows[0])
        except IndexError:
            self.rows.set("N/A")
        self.rows.event_generate("<<ComboboxSelected>>")

    def row_upd(self, event):
        # Update gui_table and total calculation based on row selection.
        self.row_val = self.rows.get()
        if self.row_val == "N/A":
            view_products(self.sql_table, "All", "All", self.gui_table)
        else:
            view_products(self.sql_table, self.column_val, self.row_val,
                          self.gui_table)
        self.total_calc()

    def total_calc(self):
        # Returns the sum of all values in a table's chosen column.
        # Used by the View_Frame class to display output.
        self.total = 0
        self.conn = dbx_aws.db_conn()
        if self.sql_table == "barrels":
            try:
                self.pg_ind = self.gui_table.columns.index("Proof Gallons")
                self.type_ind = self.gui_table.columns.index("Type")
                self.query = "SELECT total_per_pg FROM estimated_cogs"
                self.cogs = dbx_aws.db_exec(self.conn, self.query).data
                self.rum_cogs = self.cogs[0][0]
                self.whsk_cogs = self.cogs[1][0]
                for child in self.gui_table.get_children():
                    if (self.gui_table.item(child)["values"][self.type_ind]
                        == "Rum"):
                        self.total += (
                  Decimal(self.gui_table.item(child)["values"][self.pg_ind])
                  * Decimal(self.rum_cogs)
                                      )
                    else:
                        self.total += (
                  Decimal(self.gui_table.item(child)["values"][self.pg_ind])
                  * Decimal(self.whsk_cogs)
                                      )
            except:
                pass
        else:
            try:
                self.price_ind = self.gui_table.columns.index("Price")
                self.amount_ind = self.gui_table.columns.index("Amount")
                for child in self.gui_table.get_children():
                    self.total += (
              Decimal(self.gui_table.item(child)["values"][self.price_ind][1:])
              * Decimal(self.gui_table.item(child)["values"][self.amount_ind])
                                  )
            except:
                try:
                    self.price_ind = self.gui_table.columns.index("Price")
                    for child in self.gui_table.get_children():
                        self.total += (
              Decimal(self.gui_table.item(child)["values"][self.price_ind][1:])
                                      )
                except:
                    pass
        try:
            self.text_var.set("${:,.2f}".format(round(self.total, 2)))
        except:
            self.text_var.set("$0.00")

class BarrelCountFrame(tk.LabelFrame):
    # Frame containing the number of empty barrels, their price and the
    # total price. Also contains buttons to update these values.
    def __init__(self, master):
        self.master = master
        tk.LabelFrame.__init__(self, master=self.master,
                               text="Barrel Valuation", relief='ridge',
                               font="bold", bd=5, padx=2, pady=2)

        (tk.Label(
            self, text="Full Barrels", bg="dark slate gray", fg="white"
                 ).grid(row=0, column=0, sticky="NESW")
        )
        tk.Entry(self, justify='center').grid(row=1, column=0)
        (tk.Label(
            self, text="Empty Barrels", bg="dark slate gray", fg="white"
                 ).grid(row=0, column=1, sticky="NESW")
        )
        (tk.Entry(
            self, justify='center', validate='key',
            validatecommand=(self.register(valid_dig), '%S', '%d')
                 ).grid(row=1, column=1)
        )
        (tk.Label(
            self, text="Price ($)", bg="dark slate gray", fg="white"
                 ).grid(row=2, column=0, sticky="NESW")
        )
        (tk.Entry(
            self, justify='center', validate='key',
            validatecommand=(self.register(valid_dec), '%S', '%s', '%d' )
                 ).grid(row=3, column=0)
        )
        (tk.Label(
            self, text="Total ($)", bg="dark slate gray", fg="white"
                 ).grid(row=2, column=1, sticky="NESW")
        )
        tk.Entry(self, justify='center').grid(row=3, column=1)

        # Checkbox used to lock/unlock entries
        self.lock_var = tk.IntVar()
        self.lock_var.set(1)
        self.lock_cbox = tk.Checkbutton(
            self, text="Lock/Unlock",
            variable=self.lock_var,
            command=self.cbox_check)
        self.lock_cbox.grid(row=4, column=0)

        (HoverButton(
            self, text="Update", font='Arial 9 bold',
            command=self.barr_update
                    ).grid(row=4, column=1, sticky="NESW")
        )
        self.barr_update(first=1)
        self.pack(anchor='center')

    def cbox_check(self):
        # Checkbox selection function to lock/unlock entries
        if self.lock_var.get() == 1:
            for entry in self.entries[1:3]: # empty barrels/price entry
                entry.config(state='readonly')
        else:
            for entry in self.entries[1:3]: # empty barrels/price entry
                entry.config(state='normal')

    def barr_update(self, first=0):
        # Update barr_count_fr widget to have updated table values
        self.conn = dbx_aws.db_conn()
        self.query = "SELECT * FROM barrel_count"
        self.barr_count_vals = dbx_aws.db_exec(self.conn, self.query).data[0]
        self.barr_count_vals = list(self.barr_count_vals)
        self.entries = [x for x
                        in reversed(self.grid_slaves())
                        if x.winfo_class() == 'Entry']
        if first == 1:  # First time running function
            self.barr_upd_vals = self.barr_count_vals
        else:
            self.entries_vals = [x.get() for x in self.entries]
            self.sql_edit = tuple(self.entries_vals + self.barr_count_vals)
            edit_db(self.sql_edit, 'barrel_count', None, None)
            self.conn = dbx_aws.db_conn()
            self.query = "SELECT * FROM barrel_count"
            self.barr_upd_vals = dbx_aws.db_exec(self.conn, self.query).data[0]
            self.barr_upd_vals = list(self.barr_upd_vals)
        for (entry, val) in zip(self.entries, self.barr_upd_vals):
            entry.config(state='normal')
            entry.delete(0, 'end')
            entry.insert(0, round(val, 2))
            entry.config(state='readonly')
            self.lock_var.set(1)


class CogsView(tk.Toplevel):
    # Toplevel to view and edit estimated_cogs table.
    def __init__(self, master, sql_table, gui_table, view_fr):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.view_fr = view_fr
        self.x = x + 100
        self.y = y + 100
        tk.Toplevel.__init__(self, master=self.master)

        self.whiskey_fr = tk.LabelFrame(self, text="Whiskey COGS",
                                        font="Arial 12 bold")
        self.rum_fr = tk.LabelFrame(self, text="Rum COGS", font="Arial 12 bold")
        self.conn = dbx_aws.db_conn()
        self.query = "SELECT * FROM estimated_cogs"
        self.cogs_values = dbx_aws.db_exec(self.conn, self.query).data
        self.cogs_labels = [
            "Raw Materials", "Energy", "Labor",
            "Error", "Total Per Bottle", "Bond Ins.",
            "Storage", "Mult Factor", "Total Per PG"
        ]

        for (frame, l) in zip([self.whiskey_fr, self.rum_fr],
                               self.cogs_values):
            for (ind, item, desc) in zip(range(len(l)),
                                         l,
                                         self.cogs_labels):
                tk.Label(frame, text=desc).grid(row=ind, column=0)
                self.ent = tk.Entry(frame, justify="center")
                self.ent.insert(0, item)
                if desc in ["Total Per Bottle", "Total Per PG"]:
                    self.ent.config(state="readonly")
                else:
                    self.ent.config(
                        validate='key',
                        validatecommand=(
                            self.register(valid_dec), '%S', '%s', '%d'
                                        )
                    )
                self.ent.grid(row=ind, column=1)
        self.whiskey_fr.grid(row=0, column=0, padx=5)
        self.rum_fr.grid(row=0, column=1, padx=5)

        self.button_fr = tk.Frame(self, pady=5)
        (HoverButton(
            self.button_fr, text="Update", command=self.update
                    ).grid(row=0, column=0, padx=5)
        )
        (HoverButton(
            self.button_fr, text="Cancel",
            command=lambda: self.destroy()
                    ).grid(row=0, column=1, padx=5)
        )
        self.button_fr.grid(row=1, column=0, columnspan=2)

        self.total_after()
        self.title("COGS")
        self.geometry("%dx%d+%d+%d" % (450, 247, self.x, self.y))
        self.resizable(0,0)
        self.focus()

    def update(self):
        # Update estimated_cogs table with newly inputted values.
        self.conf_ques = messagebox.askquestion("COGS Confirmation",
                            "Are you sure you want to confirm? Confirming will"+
                            " update the Cost of Goods values with your " +
                            "changes.", parent=self)
        if self.conf_ques == "no":
            return self.total_after()
        else:
            self.after_cancel(self.after_func)
            self.whsk_entries = [x.get() for x
                                 in reversed(self.whiskey_fr
                                             .grid_slaves(column=1))]
            self.rum_entries = [x.get() for x
                                in reversed(self.rum_fr.grid_slaves(column=1))]
            self.conn = dbx_aws.db_conn()
            self.cur = self.conn.cursor()
            self.query = "SELECT total_per_pg FROM estimated_cogs"
            self.cur.execute(self.query)
            self.pg_ref = self.cur.fetchall()
            self.whsk_entries.append(self.pg_ref[0][0])
            self.whsk_entries = tuple(self.whsk_entries)
            self.rum_entries.append(self.pg_ref[1][0])
            self.rum_entries = tuple(self.rum_entries)
            self.query = ("UPDATE estimated_cogs SET raw_mat=%s, energy=%s, " +
                          "labor=%s, error=%s, total_per_bottle=%s, " +
                          "bond_ins=%s, storage=%s, mult_fact=%s, " +
                          "total_per_pg=%s WHERE total_per_pg=%s")
            self.cur.execute(self.query, self.whsk_entries)
            self.cur.execute(self.query, self.rum_entries)
            self.conn.commit()
            self.cur.close()
            self.conn.close()
            self.view_fr.columns.event_generate("<<ComboboxSelected>>")
            self.destroy()

    def total_after(self):
        # Auto-complete 'total per bottle' and 'total per pg' entries
        # based on other entries.
        def total_update():

            for frame in [self.whiskey_fr, self.rum_fr]:
                self.entries = [x for x
                                in reversed(frame.grid_slaves(column=1))]
                self.entries[4].config(state="normal")
                self.entries[4].delete(0, 'end')
                self.bot_total = 0
                self.entries[8].config(state="normal")
                self.entries[8].delete(0, 'end')
                self.pg_total = 0
                for entry in self.entries[:4]:
                    try:
                        self.bot_total += Decimal(entry.get())
                    except:
                        pass
                try:
                    self.entries[4].insert(0, round(self.bot_total, 2))
                    self.entries[4].config(state="readonly")
                except:
                    pass
                try:
                    self.pg_total = ((self.bot_total
                                      * Decimal(self.entries[7].get()))
                                      + Decimal(self.entries[5].get())
                                      + Decimal(self.entries[6].get())
                    )
                except:
                    pass
                try:
                    self.entries[8].insert(0, round(self.pg_total, 2))
                    self.entries[8].config(state="readonly")
                except:
                    pass
        total_update()
        self.after_func = self.after(150, self.total_after)


class EmptrView(tk.Toplevel):
    # Toplevel used to input transaction information when inventory items
    # are taken from, or returned to Montgomery.
    def __init__(self, master, sql_table, gui_table):
        self.master = master
        self.sql_table = sql_table
        self.gui_table = gui_table
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        self.window_height = 0
        tk.Toplevel.__init__(self, master=self.master)

        self.title_fr = tk.Frame(self)
        (tk.Label(
            self.title_fr, text="Employee Transaction", font="Arial 10 bold"
                 ).pack()
        )
        self.title_fr.grid(row=0, column=0, columnspan=2, pady=5)

        self.info_fr = tk.Frame(self)
        for index,desc in enumerate(self.gui_table.columns):
            (tk.Label(
                self.info_fr, text="{}:".format(desc)
                     ).grid(row=index,column=0)
            )
            if desc == "Date":
                self.date = tk.StringVar()
                self.date_entry = tk.Entry(
                    self.info_fr, state='readonly',justify='center',
                    textvariable=self.date
                )
                self.date_entry.grid(row=index, column=1)
                self.cal_photo = ImageTk.PhotoImage(cal_image)
                self.cal_link = HoverButton(
                    self.info_fr, image=self.cal_photo,
                    command=lambda:
                    cal_button(self, self.date_entry, self.master)
                )
                self.cal_link.image = self.cal_photo
                self.cal_link.grid(row=index, column=2)
            elif desc == "Product":
                self.conn = dbx_aws.db_conn()
                self.query = "SELECT product FROM bottles"
                self.product_rows = dbx_aws.db_exec(self.conn, self.query).data
                self.product_rows = [x[0] for x in self.product_rows]
                self.products = ttk.Combobox(self.info_fr,
                                             values=self.product_rows)
                self.products.config(width=16, background="white",
                                     justify='center', state='readonly')
                self.products.set(self.product_rows[0])
                self.products.grid(row=index, column=1)

            elif desc == "Unit":
                self.units = ttk.Combobox(self.info_fr,
                                          values=['Cases', 'Bottles'])
                self.units.config(width=16, background="white",
                                  justify='center', state='readonly')
                self.units.set('Cases')
                self.units.grid(row=index, column=1)
            elif desc == "Destination":
                self.dest_entry = tk.Entry(self.info_fr)
                self.dest_entry.grid(row=index, column=1)
            else:
                tk.Entry(self.info_fr).grid(row=index,column=1)
            self.window_height += 35
        self.info_fr.grid(row=1,column=0,columnspan=2)
        # Checkbox used indicate Montgomery return
        self.check_var = tk.IntVar()
        self.check_var.set(1)
        self.check_b = tk.Checkbutton(
            self, text="Items were returned to Montgomery",
            variable=self.check_var, command=self.cbox_check
        )
        self.check_b.grid(row=2, column=0, columnspan=2)
        self.cbox_check()

        self.button_fr = tk.Frame(self)
        (HoverButton(
            self.button_fr, text="Confirm", width=10,
            command=lambda: self.confirm()
                    ).pack(side='left', padx=5, pady=5)
        )
        (HoverButton(
            self.button_fr, text="Cancel", width=10,
            command=lambda: self.destroy()
                    ).pack(side='left', padx=5, pady=5)
        )
        self.button_fr.grid(row=3, column=0, columnspan=2)

        self.title("Employee Transasction")
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        # Update 'employee_transactions' table with entry information.
        # Also update 'bottles'/'samples' inventories.
        self.entries = [x.get() for x
                        in reversed(self.info_fr.grid_slaves())
                        if x.winfo_class() == "Entry"
                        or x.winfo_class() == "TCombobox"]
        if not all(self.entries):
            messagebox.showerror("Input Error",
                                 "Please make sure all of the entries are " +
                                 "completed and then try again.", parent=self)
            return
        if self.entries[3] == "Cases":
            self.inv_tbl = 'bottles'
        else:
            self.inv_tbl = 'samples'
        self.conn = dbx_aws.db_conn()
        self.cur = self.conn.cursor()
        self.query = ("INSERT INTO employee_transactions " +
                      "VALUES (%s,%s,%s,%s,%s,%s)")
        self.cur.execute(self.query, tuple(self.entries))
        if self.check_var.get() == 1:
            self.operator = "+"
        else:
            self.operator = "-"
        self.query = ("UPDATE {} SET amount=(amount {} %s) " +
                      "WHERE product=%s").format(self.inv_tbl, self.operator)
        self.cur.execute(self.query, (self.entries[2], self.entries[1]))
        self.conn.commit()
        self.cur.close()
        self.conn.close()
        db_update()
        view_products('employee_transactions', 'All', 'All', self.gui_table)
        self.destroy()

    def cbox_check(self):
        # Checkbox selection function
        if self.check_var.get() == 1:
            self.dest_entry.delete(0, 'end')
            self.dest_entry.insert(0, "Montgomery")
            self.dest_entry.config(state='readonly')
        else:
            self.dest_entry.config(state="normal")
            self.dest_entry.delete(0, 'end')


class EmptyBarrelView(tk.Toplevel):
    # Toplevel used to input date for when selected barrel was emptied.
    # Removes barrel from 'barrels' and places in 'empty_barrels'.
    def __init__(self, master, barrel_info):
        self.master = master
        self.barrel_info = barrel_info
        self.x = (screen_width/2) - (width/2) + 100
        self.y = ((screen_height/2) - (height/2)) + 50
        tk.Toplevel.__init__(self, master=self.master)

        self.main_fr = tk.LabelFrame(
            self, text="Empty Barrel: {}".format(self.barrel_info[0]),
            font="Arial 10 bold"
        )
        tk.Label(self.main_fr, text="Empty Date:").grid(row=0, column=0)
        self.date_entry = tk.Entry(
            self.main_fr, justify='center', state='readonly'
        )
        self.date_entry.grid(row=0, column=1)
        self.cal_link = HoverButton(
            self.main_fr, image=cal_photo,
            command=lambda:
            cal_button(self, self.date_entry, self.master)
        )
        self.cal_link.image = cal_photo
        self.cal_link.grid(row=0,column=3)
        tk.Label(self.main_fr, text="Remaining PG:").grid(row=1, column=0)
        (tk.Entry(
            self.main_fr, validate='key',
            validatecommand=(self.register(valid_dec), '%S', '%s', '%d')
                 ).grid(row=1, column=1)
        )
        self.add_empty = tk.IntVar()
        self.add_empty.set(1)
        self.add_empty_b = tk.Checkbutton(
            self.main_fr, text="Add empty barrel to inventory?",
            variable=self.add_empty
        )
        self.add_empty_b.grid(row=2, column=0, columnspan=2)
        self.main_fr.grid(row=0, column=0, pady=5, padx=10, ipadx=2, ipady=2)

        self.button_fr = tk.Frame(self)
        (HoverButton(
            self.button_fr, text="Confirm", command=self.confirm
                    ).pack(side='left', padx=5)
        )
        HoverButton(self.button_fr, text="Cancel").pack(side='left', padx=5)
        self.button_fr.grid(row=1, column=0, pady=5)

        self.title("Empty Barrel: {}".format(self.barrel_info[0]))
        self.focus()
        self.geometry("+%d+%d" % (self.x, self.y))
        self.resizable(0,0)

    def confirm(self):
        self.conf_quest = messagebox.askquestion(
            "Empty Barrel " + self.barrel_info[0] + "?",
            "Are you sure you want to confirm? Make sure everything is " +
            "entered correctly before continuing.", parent=self)
        if self.conf_quest == "yes":
            self.entry_vals = [x.get() for x
                               in reversed(self.main_fr.grid_slaves())
                               if x.winfo_class() == 'Entry']
            if all(self.entry_vals):
                self.conn = dbx_aws.db_conn()
                self.cur = self.conn.cursor()
                self.query = "DELETE FROM barrels WHERE barrel_no=%s"
                self.cur.execute(self.query, (self.barrel_info[0]))
                self.final_age = (
                    datetime.strptime(self.entry_vals[0], "%Y-%m-%d")
                    - datetime.strptime(self.barrel_info[4], "%Y-%m-%d")
                ).days
                self.ins_values = (self.barrel_info[:4] +
                                   [self.entry_vals[1]] +
                                   [self.barrel_info[4]] +
                                   [self.entry_vals[0]] +
                                   [self.final_age] +
                                   [self.barrel_info[6]]
                )
                self.ins_values = tuple(self.ins_values)
                self.query = ("INSERT INTO empty_barrels " +
                              "VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)")
                self.cur.execute(self.query, self.ins_values)
                if self.add_empty.get() == 1:
                    self.query = ("UPDATE barrel_count " +
                                  "SET empty_amount=(empty_amount + 1)")
                    self.cur.execute(self.query)
                self.conn.commit()
                self.cur.close()
                self.conn.close()
            else:
                messagebox.showerror(
                    "Input Error",
                    "Please make sure all of the entries have values.",
                    parent=self)
        else:
            return
        self.destroy()
        db_update()
        try:
            barr_vfr.columns.set("All")
            barr_vfr.columns.event_generate("<<ComboboxSelected>>")
        except:
            pass
        barr_count_fr.barr_update(first=1)

class OptionFrame(tk.LabelFrame):
    # Frame used to place buttons with inventory functionality.
    def __init__(self, master):
        self.master = master
        self.height = height
        tk.LabelFrame.__init__(self, master=self.master, text="Options",
                               height=self.height, relief='ridge', font="bold",
                               bd=5)


class LogisticsButton(HoverButton):
    # Command frame button
    def __init__(self, master, text, sqlite_table, gui_table, command):
        self.sqlite_table = sqlite_table
        self.gui_table = gui_table
        HoverButton.__init__(self, master=master, text=text, width=20, height=1,
                             font=('Calibri', 12, 'bold'), command=command)
        self.pack(anchor='center')


class TreeviewTable(ttk.Treeview):
    # Creates a gui_table with given columns.
    # Has ability to be sorted by gui_table_sort when column headers are
    # clicked upon.
    def __init__(self, master, columns):
        self.columns = columns
        self.width = int(table_width / (len(self.columns)))
        ttk.Treeview.__init__(self, master, columns=self.columns,
                              show='headings', height=600,
                              style="Custom.Treeview")

        if 'Date Filled' in self.columns:
            self.widths_for_cols = [75, 65, 65, 95, 108, 240, 175]
        else:
            self.widths_for_cols = [self.width] * len(self.columns)
        for i in range(len(self.columns)):
            self.column(self.columns[i], anchor='center', width=self.widths_for_cols[i])
            self.heading(str('#' + str((i+1))),
                         text=self.columns[i],
                         command=lambda col=self.columns[i]:
                         gui_table_sort(self, col, False))


        self.y_scrollbar = ttk.Scrollbar(self, orient='vertical',
                                         command=self.yview)
        self.y_scrollbar.pack(side='right', fill='y')
        self.configure(yscrollcommand=self.y_scrollbar.set)
        self.pack(side='right', fill='both', expand=1)




# Used to search for the string literal within a filename that occurs
# before the file extension (Ex. '.txt').
fileRegex = re.compile(r'''
                       ([a-zA-Z0-9_ -]+)
                       (.)
                       ([a-zA-Z_0-9])''',re.VERBOSE)

# Used to search for the the po-number in purchase order file names.
poRegex = re.compile(r'''
                     ([a-zA-Z0-9_]+)
                     (-)
                     ([0-9]{3})''',re.VERBOSE)

# Used to find important parts of mash number strings.
mashRegex = re.compile(r'''
                       (^\d{4})
                       (/)
                       (\d{2})
                       (-)
                       (\d{1})
                       ([a-zA-Z]{1})''',re.VERBOSE)

def valid_dig(str, act):
    # Entry validation used to ensure only digits.
    if act == '0':
        return True
    else:
        return str.isdigit()


def valid_dec(str, cur_str, act):
    # Entry validation used to ensure only decimal numbers.
    if act == '0':
        return True
    elif str == "." and cur_str.find(".") != -1:
        return False
    elif str == "." and cur_str.find(".") == -1:
        return True
    elif str.count(".") <= 1:
        return True
    else:
        return str.isdigit()


def disable_event():
    # Prevent user from 'x'-ing out of sample-desc entry to
    # prevent issue with confirm function.
    pass

def busy_cursor(self, function):
    # Display a busy cursor while the function executes
    # People love visual feedback!
    try:
        if self.total_after:
            self.after_cancel(self.total_after)
        self.master.config(cursor='wait')
        self.config(cursor='wait')
        self.master.update()
        threading.Thread(target=function).start()
    except: # Prevents cursor from permantently becoming 'wait'
        pass


# Option values for dropdown menus.
type_options = {
    'raw_materials' : ['Bottles', 'Boxes', 'Caps', 'Capsules', 'Labels'],
    'bottles' : ['Vodka', 'Whiskey', 'Rum', 'Seltzer', 'Other'],
    'barrels' : ['Bourbon', 'Rye', 'Malt', 'Rum', 'Other'],
    'grain' : ['Corn', 'Rye', 'Malted Barley', 'Malted Wheat', 'Wheat', 'Oat',
               'Molasses'],
    'samples' : ['Vodka', 'Whiskey', 'Rum', 'Seltzer', 'Other'],
    'mashes' : ['Bourbon', 'Rye', 'Malt', 'Rum', 'Other'],
    'grain_log' : ['Corn', 'Rye', 'Malted Barley', 'Malted Wheat', 'Oat',
                   'Molasses']}
